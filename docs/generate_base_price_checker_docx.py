"""
Generates: Tri-E Enterprises - Base Price Checker User Guide.docx

A step-by-step user guide for staff on recording each supplier's base price
for a product, and using the side-by-side comparison shown on Purchase
Orders to check prices before confirming a line's cost.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

ASSETS = os.path.join(os.path.dirname(__file__), '_baseprice_assets')

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Default paragraph style ───────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ── Heading colours ───────────────────────────────────────────────────────────
HEADING_COLORS = {
    1: RGBColor(0x1E, 0x3A, 0x5F),   # dark navy
    2: RGBColor(0x2E, 0x6D, 0xA4),   # medium blue
    3: RGBColor(0x2E, 0x86, 0xAB),   # steel blue
    4: RGBColor(0x45, 0x8B, 0x74),   # teal green
}

def set_heading(text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.color.rgb = HEADING_COLORS.get(level, RGBColor(0, 0, 0))
    run.font.bold = True
    run.font.size = Pt({1: 18, 2: 14, 3: 12, 4: 11}.get(level, 11))
    return h

def shade_cell(cell, hex_color='D9EAF7'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], 'D9EAF7')
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                shade_cell(cells[ci], 'F5FAFF')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return t

def add_note(text, color=RGBColor(0x44, 0x6E, 0x91), icon='ℹ'):
    p = doc.add_paragraph()
    run = p.add_run(icon + '  ' + text)
    run.font.italic = True
    run.font.size   = Pt(10)
    run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.2)
    return p

def add_warning(text):
    return add_note(text, color=RGBColor(0xA6, 0x3D, 0x2E), icon='⚠')

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    text = re.sub(r'^[-*]\s*', '', text)
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    return p

def add_para(text):
    p = doc.add_paragraph()
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    return p

def add_image(filename, width_in=5.5, caption=None):
    path = os.path.join(ASSETS, filename)
    if not os.path.exists(path):
        add_warning(f'Screenshot not found: {filename}')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cr = cap.add_run(caption)
        cr.font.name = 'Courier New'
        cr.font.size = Pt(8.5)
        cr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        cap.paragraph_format.space_after = Pt(10)

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Base Price Checker', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
title.runs[0].font.size = Pt(26)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('User Guide')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = meta.add_run('Tri-E Enterprises · TOS System · Purchases &amp; Products Module')
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = meta2.add_run('Prepared: August 26, 2026')
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
set_heading('Overview', 1)
add_para('Many products can be bought from more than one supplier, often at different prices — like the JBJ, Casini, and Fussen price list you already keep for steel products. The **Base Price Checker** brings that comparison into the system: record what each supplier charges for a product once, and every time you create a Purchase Order for that product, you\'ll see all recorded supplier prices side by side.')
doc.add_paragraph()
add_para('This is a **comparison tool, not an automatic price-setter** — it never changes the price you enter on a Purchase Order. It just shows you what each supplier is known to charge, so you can catch an overpriced quote before confirming it.')
doc.add_paragraph()
add_note('There are two steps: (1) record supplier base prices on the **Product**, and (2) view the comparison while creating a **Purchase Order**.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 1. RECORDING SUPPLIER BASE PRICES
# ════════════════════════════════════════════════════════════════════════════
set_heading('1. Recording Supplier Base Prices', 1)
add_para('Do this once per product — whenever you learn a new supplier\'s price, or a price changes.')
doc.add_paragraph()

add_numbered('Go to **Inventory & Sales → Products** and open the product you want to price-check (or create a new one).')
add_numbered('Scroll to the **Supplier Base Prices** section.')
add_numbered('Click **Add to supplier Base Prices**.')
add_numbered('Choose a **Supplier** and enter their **Base Price** for this product.')
add_numbered('Repeat for every supplier you buy this product from.')
add_numbered('Click **Save changes**.')
doc.add_paragraph()

add_image('product_supplier_prices.png', width_in=5.3,
          caption='Product → Edit → Supplier Base Prices (example: Angle Bar 3mm x 1 (1/4))')

add_note('You can add as many suppliers as you want, and update a price at any time — just edit the amount and save again.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 2. COMPARING PRICES ON A PURCHASE ORDER
# ════════════════════════════════════════════════════════════════════════════
set_heading('2. Comparing Prices on a Purchase Order', 1)
add_para('Once a product has base prices recorded, the comparison shows up automatically while you\'re building a Purchase Order.')
doc.add_paragraph()

add_numbered('Go to **Inventory & Sales → Purchases** and click **Create** (or open an existing purchase and click **Edit**).')
add_numbered('Choose the **Supplier** you\'re ordering from.')
add_numbered('In the **Purchase Items** section, choose the **Product** on the line item.')
add_numbered('A **Base Price by Supplier** list appears, showing every supplier\'s recorded price for that product — cheapest first.')
add_numbered('The supplier you chose in step 2 is shown in **bold**, marked "(selected supplier)", so you can immediately see if their price is the best one available.')
doc.add_paragraph()

add_image('purchase_price_comparison.png', width_in=5.3,
          caption='Purchase Order line item — Base Price by Supplier comparison')

add_para('In this example, the Purchase Order is being placed with **Casini**, and the comparison shows Casini\'s price (₱332.00) is actually higher than JBJ\'s (₱320.00) for the same product — a signal worth double-checking before confirming the order.')
doc.add_paragraph()

add_note('If a product has no supplier base prices recorded yet, the panel simply says "No supplier base prices recorded yet for this product." — nothing else changes, and you can still fill in the Purchase Order as normal.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 3. GOOD TO KNOW
# ════════════════════════════════════════════════════════════════════════════
set_heading('3. Good to Know', 1)
add_bullet('**It\'s read-only.** The comparison never fills in or overrides the Price field on the Purchase Order — you still type in the actual cost yourself.')
add_bullet('**It\'s per product.** Each product keeps its own list of supplier prices; there\'s nothing to set up per Purchase Order beyond picking the product.')
add_bullet('**It updates instantly.** Changing the product or the supplier on the Purchase Order refreshes the comparison right away — no need to save first.')
add_bullet('**Prices can be edited anytime.** If a supplier changes their price, just update it on the Product\'s Supplier Base Prices section — every future Purchase Order will reflect the new number.')
doc.add_paragraph()

foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot.add_run('Tri-E Enterprises · TOS System')
fr.font.size = Pt(8.5)
fr.italic = True
fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

OUT_PATH = os.path.join(os.path.dirname(__file__), 'Tri-E Enterprises - Base Price Checker User Guide.docx')
doc.save(OUT_PATH)
print('Saved:', OUT_PATH)
