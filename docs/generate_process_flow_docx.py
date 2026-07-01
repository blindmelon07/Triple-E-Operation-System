"""
Generates docs/system-process-flow.docx — a full process-flow document
for the TOS system, covering every major business process end to end.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Default paragraph style ───────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
BLUE   = RGBColor(0x2E, 0x6D, 0xA4)
STEEL  = RGBColor(0x2E, 0x86, 0xAB)
TEAL   = RGBColor(0x45, 0x8B, 0x74)
GRAY   = RGBColor(0x55, 0x55, 0x55)

BOX_START  = 'C6E0B4'   # green  — start / end
BOX_PROC   = 'D9EAF7'   # blue   — process step
BOX_DEC    = 'FFE699'   # amber  — decision
BOX_DATA   = 'E2E2E2'   # gray   — system/data event

HEADING_COLORS = {1: NAVY, 2: BLUE, 3: STEEL, 4: TEAL}


def set_heading(text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.color.rgb = HEADING_COLORS.get(level, RGBColor(0, 0, 0))
    run.font.bold = True
    run.font.size = Pt({1: 20, 2: 15, 3: 12.5, 4: 11}.get(level, 11))
    return h


def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_border_color(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '8')
        el.set(qn('w:color'), hex_color)
        borders.append(el)
    tcPr.append(borders)


def add_para(text, bold=False, italic=False, size=10.5, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = bold
    r.font.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], 'D9EAF7')
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                shade_cell(cells[ci], 'F5FAFF')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return t


# ── Flow-diagram primitives ───────────────────────────────────────────────────
# A "flow" is a stack of single-cell boxes connected by centred down-arrows.
# Box kinds: start, end, process, decision, data

def _box(text, fill, text_color=RGBColor(0x20, 0x20, 0x20), width=5.5, bold=False):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.width = Inches(width)
    shade_cell(cell, fill)
    set_cell_border_color(cell, '8C8C8C')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.bold = bold
    r.font.color.rgb = text_color
    for row in t.rows:
        row.cells[0].width = Inches(width)
    return t


def _arrow(label=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('↓' + (f'  {label}' if label else ''))
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = GRAY
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    return p


def add_flow(steps):
    """
    steps: list of dicts:
      {'type': 'start'|'end', 'text': ...}
      {'type': 'process', 'text': ...}
      {'type': 'data', 'text': ...}
      {'type': 'decision', 'text': ..., 'yes': ..., 'no': ...}
    """
    for i, step in enumerate(steps):
        kind = step['type']
        if kind in ('start', 'end'):
            _box(step['text'], BOX_START, bold=True)
        elif kind == 'process':
            _box(step['text'], BOX_PROC)
        elif kind == 'data':
            _box(step['text'], BOX_DATA, text_color=RGBColor(0x40, 0x40, 0x40))
        elif kind == 'decision':
            _box('◆ ' + step['text'], BOX_DEC, bold=True)
            # Yes / No outcome row
            t = doc.add_table(rows=1, cols=2)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            left, right = t.rows[0].cells
            left.width = Inches(2.7)
            right.width = Inches(2.7)
            shade_cell(left, 'E9F5E1')
            shade_cell(right, 'FCE4E4')
            set_cell_border_color(left, '8C8C8C')
            set_cell_border_color(right, '8C8C8C')
            lp = left.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lr1 = lp.add_run('YES → ')
            lr1.font.bold = True
            lr1.font.size = Pt(9.5)
            lr1.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            lr2 = lp.add_run(step['yes'])
            lr2.font.size = Pt(9.5)
            rp = right.paragraphs[0]
            rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr1 = rp.add_run('NO → ')
            rr1.font.bold = True
            rr1.font.size = Pt(9.5)
            rr1.font.color.rgb = RGBColor(0xB0, 0x2A, 0x2A)
            rr2 = rp.add_run(step['no'])
            rr2.font.size = Pt(9.5)
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
        if i < len(steps) - 1:
            _arrow()
    doc.add_paragraph()


def add_detail_list(details):
    """details: list of (bold_lead, rest_text) tuples rendered as a numbered list."""
    for lead, rest in details:
        p = doc.add_paragraph(style='List Number')
        r1 = p.add_run(lead)
        r1.font.bold = True
        r1.font.size = Pt(10.5)
        if rest:
            r2 = p.add_run('  ' + rest)
            r2.font.size = Pt(10.5)
    doc.add_paragraph()


def add_legend():
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = [('Start / End', BOX_START), ('Process Step', BOX_PROC),
              ('Decision', BOX_DEC), ('System / Data Event', BOX_DATA)]
    for i, (label, color) in enumerate(labels):
        cell = t.rows[0].cells[i]
        shade_cell(cell, color)
        set_cell_border_color(cell, '8C8C8C')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.font.size = Pt(9)
        r.font.bold = True
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('TOS — System Process Flow', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = NAVY
title.runs[0].font.size = Pt(26)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Tri-E Enterprises — Business Process Flow Documentation')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = meta.add_run('Covers: Point of Sale · Quotations · Purchasing · Delivery · Void/Refunds ·\n'
                   'Cash Register · Fleet Maintenance · Attendance & Leave · Payroll · Access & Audit')
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual list)
# ════════════════════════════════════════════════════════════════════════════
set_heading('Contents', 1)
toc_items = [
    '1. System Overview & Actors',
    '2. Point-of-Sale (POS) Sales Process',
    '3. Quotation Approval & Conversion Process',
    '4. Purchasing / Procurement Process',
    '5. Delivery Process',
    '6. Void / Refund Process',
    '7. Cash Register Session Process',
    '8. Vehicle Maintenance Process',
    '9. Attendance & Leave Process',
    '10. Payroll Process',
    '11. User Access, Roles & Audit Process',
    '12. Consolidated End-to-End System Flow',
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW & ACTORS
# ════════════════════════════════════════════════════════════════════════════
set_heading('1. System Overview & Actors', 1)
add_para(
    'TOS is a Laravel + Filament application combining a browser-based Point-of-Sale (POS) '
    'terminal with a full back-office admin panel covering inventory, purchasing, delivery, '
    'finance, fleet maintenance, HR, and payroll. The sections below describe how a '
    'transaction or request moves through the system from start to finish, who is involved '
    'at each step, and what the system does automatically in the background.'
)

set_heading('1.1 Key Actors', 2)
add_table(
    ['Actor', 'Role in the System'],
    [
        ['Cashier', 'Operates the POS terminal — opens/closes the register, builds sales, applies discounts, takes payment, requests voids.'],
        ['Admin / Manager', 'Approves quotations, void requests, service requests, leave requests, and payroll; manages master data and reports.'],
        ['Driver', 'Picks up and delivers orders assigned through the Delivery Management module.'],
        ['Employee', 'Logs attendance, files leave requests, and is paid through Payroll.'],
        ['Supplier', 'External party — fulfils Purchase orders, fully or partially.'],
        ['Customer', 'Buys products (walk-in or on file), optionally on credit terms, and may request a quotation first.'],
        ['System (automated)', 'Adjusts inventory, cash register totals, and the Audit Log automatically as a side effect of the actions above.'],
    ],
    col_widths=[1.6, 4.6]
)

set_heading('1.2 Flow Diagram Legend', 2)
add_legend()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 2. POS SALES PROCESS
# ════════════════════════════════════════════════════════════════════════════
set_heading('2. Point-of-Sale (POS) Sales Process', 1)
add_para(
    'The core revenue-generating flow. A cashier must have an open cash register session '
    'before any sale can be processed.'
)

pos_flow = [
    {'type': 'start', 'text': 'Cashier logs in and opens /pos'},
    {'type': 'process', 'text': 'Open Cash Register (enter opening cash amount)'},
    {'type': 'process', 'text': 'Add products to cart\n(standard, weight/volume, or manual custom items)'},
    {'type': 'process', 'text': 'Adjust quantity, apply per-item Discount (₱, per-piece or flat-total),\nand optional Delivery Fee'},
    {'type': 'process', 'text': 'Select Customer (search existing, add new inline, or leave blank for Walk-in)'},
    {'type': 'decision', 'text': 'Save as Quotation instead of completing the sale now?',
     'yes': 'Go to Quotation Process (Section 3)', 'no': 'Continue to Payment'},
    {'type': 'process', 'text': 'Click Charge → choose Payment Method\n(Cash / Bank Transfer / Check / Credit Card / GCash / Maya)'},
    {'type': 'decision', 'text': 'Apply Payment Terms (credit sale)?',
     'yes': 'Choose 5/10/15/30/60-day term → due_date set, payment_status = unpaid, amount_paid = 0',
     'no': 'payment_status = paid, amount_paid = total'},
    {'type': 'process', 'text': 'Confirm Payment'},
    {'type': 'data', 'text': 'Sale + SaleItem records created;\ninventory decremented for non-manual items'},
    {'type': 'decision', 'text': 'Was this a paid, non-credit sale?',
     'yes': 'Cash register totals updated (total sales, cash sales, transaction count)',
     'no': 'Excluded from register totals until later collected'},
    {'type': 'process', 'text': 'Success modal → print receipt (thermal or delivery format)'},
    {'type': 'end', 'text': 'Transaction complete'},
]
add_flow(pos_flow)

set_heading('2.1 Step-by-Step Detail', 2)
add_detail_list([
    ('Open Register —', 'one open session per user at a time; the session tracks running totals until closed.'),
    ('Add to Cart —', 'weight/volume/length products (kg, g, L, mL, m, ft) prompt a quantity modal; manual items are flagged is_manual and never touch inventory.'),
    ('Discount —', 'entered per line item in ₱; by default it is per piece (× quantity), or a flat total for the line if the "flat total" checkbox is ticked. Delivery Fee is a single amount added to the whole sale.'),
    ('Customer —', 'a blank selection records the sale as Walk-in Customer.'),
    ('Payment Method —', 'six supported methods; Payment Terms can be layered on top of any method to turn it into a credit sale.'),
    ('Confirm Payment —', 'creates the Sale and its SaleItems in one transaction; if this sale originated from an approved quotation, that quotation is marked converted_to_sale.'),
    ('Register Totals —', 'only immediately-paid (non-credit) sales count toward the open register session; credit/term sales are excluded until collected.'),
    ('Receipt —', 'printed as either a thermal receipt or a delivery-style receipt with signature lines.'),
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 3. QUOTATION PROCESS
# ════════════════════════════════════════════════════════════════════════════
set_heading('3. Quotation Approval & Conversion Process', 1)
add_para('Lets a cashier offer a formal, time-bound price quote before committing to a sale.')

quote_flow = [
    {'type': 'start', 'text': 'Cashier builds a cart and opens the Quotation modal'},
    {'type': 'process', 'text': 'Add notes and set validity period (default 30 days)'},
    {'type': 'process', 'text': 'Save Quotation'},
    {'type': 'data', 'text': 'Quotation record created (status = Pending); QuotationItems saved;\ninventory NOT affected'},
    {'type': 'process', 'text': 'Email notification sent to admins with approve_quotation permission'},
    {'type': 'decision', 'text': 'Admin reviews the quotation',
     'yes': 'status = Approved', 'no': 'status = Rejected (process ends)'},
    {'type': 'process', 'text': 'Staff opens POS with the quotation ID → cart pre-filled with items,\nquantities, and discounts'},
    {'type': 'process', 'text': 'Complete the sale normally (Section 2)'},
    {'type': 'data', 'text': 'Quotation status → Converted to Sale'},
    {'type': 'end', 'text': 'Quotation closed'},
]
add_flow(quote_flow)

add_para('A quotation left un-converted past its valid_until date is treated as Expired.', italic=True, size=9.5, color=GRAY)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 4. PURCHASING PROCESS
# ════════════════════════════════════════════════════════════════════════════
set_heading('4. Purchasing / Procurement Process', 1)
add_para('Records stock ordered from and received from suppliers.')

purchase_flow = [
    {'type': 'start', 'text': 'Staff creates a Purchase for a Supplier'},
    {'type': 'process', 'text': 'Add Purchase Items (product, quantity ordered, unit price)'},
    {'type': 'process', 'text': 'Supplier delivers stock'},
    {'type': 'decision', 'text': 'Received in full?',
     'yes': 'quantity_received = quantity ordered for every item → status = Received',
     'no': 'quantity_received partially entered → status = Partial (or Pending if none received yet)'},
    {'type': 'data', 'text': 'Inventory increases by quantity_received\n(create: += ; edit: adjusted by the delta; delete: -=)'},
    {'type': 'data', 'text': 'Purchase total recalculated as Σ(price × quantity ordered)'},
    {'type': 'process', 'text': 'Record payment to supplier (full or partial)'},
    {'type': 'data', 'text': 'payment_status updates (unpaid / partial / paid);\ndue_date follows the supplier’s payment terms'},
    {'type': 'end', 'text': 'Purchase closed once fully paid'},
]
add_flow(purchase_flow)

add_para(
    'The purchase total reflects the ordered quantity × price (what is owed), independent of '
    'how much has physically arrived — receiving status and payable amount are tracked separately.',
    italic=True, size=9.5, color=GRAY
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 5. DELIVERY PROCESS
# ════════════════════════════════════════════════════════════════════════════
set_heading('5. Delivery Process', 1)
add_para('Tracks a completed sale from dispatch to the customer’s door.')

delivery_flow = [
    {'type': 'start', 'text': 'Sale completed and marked for delivery'},
    {'type': 'data', 'text': 'Delivery record created (status = Pending)'},
    {'type': 'process', 'text': 'Assign a Driver'},
    {'type': 'data', 'text': 'status = Assigned'},
    {'type': 'process', 'text': 'Driver picks up the goods'},
    {'type': 'data', 'text': 'status = Picked Up (picked_up_at recorded)'},
    {'type': 'process', 'text': 'Driver travels to the customer'},
    {'type': 'data', 'text': 'status = In Transit'},
    {'type': 'decision', 'text': 'Delivered successfully?',
     'yes': 'status = Delivered (delivered_at recorded); customer rates 1–5 stars + feedback',
     'no': 'status = Failed or Returned'},
    {'type': 'end', 'text': 'Delivery closed; delivery_time_minutes calculated'},
]
add_flow(delivery_flow)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 6. VOID / REFUND PROCESS
# ════════════════════════════════════════════════════════════════════════════
set_heading('6. Void / Refund Process', 1)
add_para('Requires manager approval before any completed sale can be reversed.')

void_flow = [
    {'type': 'start', 'text': 'Cashier finds the sale in Recent Sales / Reprint and requests a Void\nwith a reason'},
    {'type': 'data', 'text': 'VoidRequest created (status = Pending)'},
    {'type': 'decision', 'text': 'Cashier cancels the request before it is reviewed?',
     'yes': 'status = Rejected ("Cancelled by cashier") — process ends',
     'no': 'Cashier polls status while waiting for a manager'},
    {'type': 'decision', 'text': 'Admin / Manager reviews the request',
     'yes': 'Approved — Sale.is_voided = true; inventory restored for non-manual items;\nregister totals reversed (if it was a paid, non-credit sale)',
     'no': 'Rejected — a rejection reason is required; the sale remains active'},
    {'type': 'end', 'text': 'Void request closed'},
]
add_flow(void_flow)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. CASH REGISTER SESSION PROCESS
# ════════════════════════════════════════════════════════════════════════════
set_heading('7. Cash Register Session Process', 1)
add_para('Wraps a cashier’s entire shift, from opening the drawer to reconciling it at close.')

register_flow = [
    {'type': 'start', 'text': 'Cashier logs into the POS'},
    {'type': 'decision', 'text': 'Does this user already have an open session?',
     'yes': 'Resume the existing session', 'no': 'Open Register — enter opening cash amount'},
    {'type': 'process', 'text': 'Process sales throughout the shift\n(Section 2)'},
    {'type': 'data', 'text': 'Each paid, non-credit sale updates total sales, cash sales, and transaction count'},
    {'type': 'process', 'text': 'Close Register — enter the physical cash count'},
    {'type': 'data', 'text': 'System calculates expected amount vs. actual count → discrepancy\n(green if over, red if short)'},
    {'type': 'data', 'text': 'PDF register closure report auto-generated and downloadable anytime after'},
    {'type': 'end', 'text': 'Session closed'},
]
add_flow(register_flow)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 8. VEHICLE MAINTENANCE PROCESS
# ════════════════════════════════════════════════════════════════════════════
set_heading('8. Vehicle Maintenance Process', 1)
add_para('Keeps the delivery fleet roadworthy through a request → approval → service workflow.')

vehicle_flow = [
    {'type': 'start', 'text': 'Staff/driver notices an issue or a scheduled service is due'},
    {'type': 'process', 'text': 'Submit a Service Request\n(vehicle, service type, priority, mileage, description, preferred date)'},
    {'type': 'data', 'text': 'status = Pending'},
    {'type': 'decision', 'text': 'Approved by manager?',
     'yes': 'estimated_cost filled in, approved_by set → status = Approved',
     'no': 'rejection_reason required → status = Rejected (process ends)'},
    {'type': 'process', 'text': 'Vehicle is serviced'},
    {'type': 'data', 'text': 'Service Record created (parts/labor cost, work performed, next service\ndate/mileage); linked back to the request; request status = Completed'},
    {'type': 'data', 'text': 'Vehicle current_mileage updated; maintenance_due re-evaluated'},
    {'type': 'end', 'text': 'Vehicle back in service'},
]
add_flow(vehicle_flow)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 9. ATTENDANCE & LEAVE PROCESS
# ════════════════════════════════════════════════════════════════════════════
set_heading('9. Attendance & Leave Process', 1)

set_heading('9.1 Daily Attendance', 2)
attendance_flow = [
    {'type': 'start', 'text': 'Employee starts their shift'},
    {'type': 'process', 'text': 'Time-in recorded'},
    {'type': 'process', 'text': 'Time-out recorded'},
    {'type': 'data', 'text': 'total_hours auto-calculated; status assigned\n(Present / Late / Half Day / Absent / On Leave)'},
    {'type': 'end', 'text': 'Attendance record saved'},
]
add_flow(attendance_flow)

set_heading('9.2 Leave Request', 2)
leave_flow = [
    {'type': 'start', 'text': 'Employee submits a Leave Request\n(leave type, start/end date, reason)'},
    {'type': 'data', 'text': 'total_days auto-calculated; status = Pending'},
    {'type': 'decision', 'text': 'Approved?',
     'yes': 'status = Approved; Attendance auto-marked "On Leave" for those dates',
     'no': 'rejection_reason required → status = Rejected'},
    {'type': 'end', 'text': 'Leave request closed'},
]
add_flow(leave_flow)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 10. PAYROLL PROCESS
# ════════════════════════════════════════════════════════════════════════════
set_heading('10. Payroll Process', 1)
add_para('Follows a strict Draft → Approved → Paid workflow; cannot be cancelled once Paid.')

payroll_flow = [
    {'type': 'start', 'text': 'Admin creates a Payroll for a pay period\n(start/end date, Daily / Weekly / Semi-Monthly)'},
    {'type': 'data', 'text': 'status = Draft'},
    {'type': 'process', 'text': 'Add Payroll Items per employee (gross pay, deductions)'},
    {'type': 'data', 'text': 'recalculateTotals(): total_gross, total_deductions, total_net computed'},
    {'type': 'decision', 'text': 'Approve payroll? (only while Draft)',
     'yes': 'status = Approved; approved_by set', 'no': 'remains Draft for further edits'},
    {'type': 'decision', 'text': 'Mark as Paid? (only while Approved)',
     'yes': 'status = Paid; paid_at recorded', 'no': 'remains Approved until payment is released'},
    {'type': 'end', 'text': 'Payroll closed (cannot be cancelled once Paid)'},
]
add_flow(payroll_flow)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 11. USER ACCESS, ROLES & AUDIT PROCESS
# ════════════════════════════════════════════════════════════════════════════
set_heading('11. User Access, Roles & Audit Process', 1)

access_flow = [
    {'type': 'start', 'text': 'Admin creates a User account and assigns one or more Roles'},
    {'type': 'data', 'text': 'Role grants resource-level permissions (view/create/update/delete)\nand custom permissions (e.g. approve_quotation)'},
    {'type': 'process', 'text': 'User logs in and performs actions across the modules above'},
    {'type': 'data', 'text': 'Every significant action (create, update, approve, void, etc.)\nis automatically written to the Audit Log'},
    {'type': 'end', 'text': 'Audit Logs remain viewable (read-only) for accountability review'},
]
add_flow(access_flow)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 12. CONSOLIDATED END-TO-END SYSTEM FLOW
# ════════════════════════════════════════════════════════════════════════════
set_heading('12. Consolidated End-to-End System Flow', 1)
add_para(
    'The diagram below ties every process together, showing how goods and money move through '
    'the system from procurement to final financial reporting.'
)

overall_flow = [
    {'type': 'start', 'text': 'Supplier stock arrives (Section 4 — Purchasing)'},
    {'type': 'data', 'text': 'Inventory increased'},
    {'type': 'process', 'text': 'Customer served at POS — direct sale or via an approved Quotation\n(Sections 2 & 3)'},
    {'type': 'data', 'text': 'Inventory decreased; Sale + payment recorded; Cash Register totals updated'},
    {'type': 'decision', 'text': 'Does the order need to be delivered?',
     'yes': 'Delivery process runs (Section 5) with a Driver', 'no': 'Customer picks up directly'},
    {'type': 'decision', 'text': 'Does the sale need to be reversed?',
     'yes': 'Void process runs (Section 6) — inventory and register totals are reversed',
     'no': 'Sale stands as final'},
    {'type': 'process', 'text': 'Register closed at end of shift (Section 7)'},
    {'type': 'data', 'text': 'Revenue, collections, COGS, and profit metrics feed the Financial\nOverview (Accounting Service)'},
    {'type': 'data', 'text': 'Fleet Maintenance (Section 8) and Payroll (Sections 9–10) costs\nare included as operating expenses'},
    {'type': 'end', 'text': 'Management reviews dashboards and Audit Logs (Section 11)'},
]
add_flow(overall_flow)

doc.add_paragraph()
add_para('— End of Document —', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)

doc.save('docs/system-process-flow.docx')
print('Saved docs/system-process-flow.docx')
