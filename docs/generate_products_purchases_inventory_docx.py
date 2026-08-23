"""
Generates: Tri-E Enterprises - Products, Purchases & Inventory In-Out Documentation.docx

Describes how the Products, Purchases, and Inventory In/Out modules work
together end-to-end, including the stock-movement logging added in the
Inventory Report fix.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Default paragraph style ───────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ── Heading colours ───────────────────────────────────────────────────────────
HEADING_COLORS = {
    1: RGBColor(0x1E, 0x3A, 0x5F),
    2: RGBColor(0x2E, 0x6D, 0xA4),
    3: RGBColor(0x2E, 0x86, 0xAB),
    4: RGBColor(0x45, 0x8B, 0x74),
}

def set_heading(text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.color.rgb = HEADING_COLORS.get(level, RGBColor(0, 0, 0))
    run.font.bold = True
    run.font.size = Pt({1: 18, 2: 14, 3: 12, 4: 11}.get(level, 11))
    return h

def shade_cell(cell, hex_color='D9EAF7'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], 'D9EAF7')
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                shade_cell(cells[ci], 'F5FAFF')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return t

def add_code_block(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F0F0F0')
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_note(text, color=RGBColor(0x44, 0x6E, 0x91), icon='ℹ'):
    p = doc.add_paragraph()
    run = p.add_run(icon + '  ' + text)
    run.font.italic = True
    run.font.size   = Pt(10)
    run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.2)
    return p

def add_warning(text):
    return add_note(text, color=RGBColor(0xA6, 0x3D, 0x2E), icon='⚠')

def add_success(text):
    return add_note(text, color=RGBColor(0x2E, 0x7D, 0x4F), icon='✔')

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    text = re.sub(r'^[-*]\s*', '', text)
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    return p

def add_para(text):
    p = doc.add_paragraph()
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    return p

def add_flow_step(number, title, desc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'  {number}  ')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    tc = p._p
    # simple bold arrow-style step (no shading engine per-run, keep it simple)
    r2 = p.add_run(f'{title} — ')
    r2.bold = True
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    r3 = p.add_run(desc)
    r3.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25)
    return p

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Products, Purchases &', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
title.runs[0].font.size = Pt(24)

title2 = doc.add_heading('Inventory In/Out', 0)
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
title2.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
title2.runs[0].font.size = Pt(24)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('How Stock Moves Through the System')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = meta.add_run('Tri-E Enterprises · TOS System · Inventory & Sales Module')
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = meta2.add_run('Prepared: August 22, 2026')
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
set_heading('1. Overview', 1)
add_para('These three areas work together to track stock from the moment it enters the store to the moment it is sold:')
doc.add_paragraph()

add_flow_step('1', 'Product', 'the catalogue entry — name, price, category, unit of measure.')
add_flow_step('2', 'Inventory', 'one live stock count per product — the number the POS and reports actually use.')
add_flow_step('3', 'Purchase', 'the record of stock ordered and received from a supplier — the main way stock goes up.')
add_flow_step('4', 'Inventory Movement', 'a permanent log entry created every time stock goes in or out, no matter the cause.')
add_flow_step('5', 'Inventory In/Out Report', 'reads the Inventory Movement log to show a complete, auditable history.')
doc.add_paragraph()

add_para('**The key relationship to understand:** a Product\'s "Stock" number and the Inventory Movement log are two different things.')
for b in [
    '**Inventory.quantity** is the single number that answers "how much do we have right now?" — it is what the POS checks before a sale and what every report displays as current stock.',
    '**Inventory Movements** are the individual "in" and "out" entries that explain *how* that number got where it is. The In/Out Report is built entirely from these entries.',
]:
    add_bullet(b)
add_note('If a stock change does not create an Inventory Movement record, the current stock number can still be correct — but the change becomes invisible in the report. This is a common source of confusion and was the cause of the issues documented separately in the Inventory Report Fix Documentation.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# PRODUCTS
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
set_heading('2. Products', 1)
add_para('Location: **Home → Products** (visible to super_admin and anyone with the Product permissions)')
add_para('Files: `ProductResource.php`, `ProductForm.php`, `ProductsTable.php`, `CreateProduct.php`, `EditProduct.php`')

set_heading('2.1 What a Product Record Holds', 2)
add_table(
    ['Field', 'Purpose'],
    [
        ['Name', 'Product name shown throughout the system'],
        ['Category', 'Used for grouping/filtering in reports'],
        ['Supplier', 'Default/primary supplier for this product'],
        ['Price', 'Selling price used at POS'],
        ['Cost Price', 'Used to calculate profit margin in reports'],
        ['Stock', 'Opening / current quantity — see 2.2 below'],
        ['Unit', 'Base unit of measure (Piece, Kilo, Liter, Bag, Roll, etc.)'],
        ['Additional Sellable Units', 'Optional alt units the product can also be sold in (e.g. sell a hose per Meter or per Roll), each with its own price and a conversion factor back to the base unit'],
    ],
    [2.2, 4.8]
)

set_heading('2.2 The "Stock" Field — Two Different Behaviors', 2)
add_para('**On Create (new product):** the Stock field sets the product\'s opening quantity. This creates the Inventory record and logs an "Initial Stock" in-movement.')
doc.add_paragraph()
add_para('**On Edit (existing product):** the Stock field is now **read-only for most roles**. Only **super_admin** and **Ops Sup** can still type into it directly. When they do, the system logs the difference as a "Manual Stock Adjustment" movement — but it is still an absolute overwrite (typing a number replaces the count, it does not add to it).')
doc.add_paragraph()
add_warning('Everyone else (Warehouse clerk, cashier, Sales Rep, Driver) sees the current stock number on Edit Product but cannot change it. They are directed to use Purchases instead, which is additive and fully audited.')

set_heading('2.3 Products List & Report', 2)
add_para('The Products table (list view) always shows **`inventory.quantity`** as "Stock" — the live, real count — not the raw product record. The **Products Report** (Reports → Products Report) adds category/supplier filtering, total units sold per product, and CSV export.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# PURCHASES
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
set_heading('3. Purchases', 1)
add_para('Location: **Home → Purchases**')
add_para('Files: `PurchaseResource.php`, `PurchaseForm.php`, `PurchasesTable.php`, `Purchase.php`, `PurchaseItem.php`')
add_para('This is the primary, intended way stock enters the system.')

set_heading('3.1 Creating a Purchase', 2)
add_numbered('Select a **Supplier**. If the supplier has payment terms configured (e.g. Net 30), the **Due Date** auto-fills from the purchase date.')
add_numbered('Set the **Payment Status** — Unpaid, Partial, or Paid — with an amount paid and paid date as needed.')
add_numbered('Add one line per product in **Purchase Items**:')
for b in [
    '**Product** — selecting it auto-fills price and unit from the product record',
    '**Ordered Qty** — how much was ordered',
    '**Received Units** — how much has actually arrived (can be less than ordered, for partial deliveries)',
    '**Price** — cost per unit for this purchase',
]:
    add_bullet(b, level=1)
add_para('The purchase **Total** is calculated automatically from all line items (price × quantity), recalculated on every change.')
doc.add_paragraph()

set_heading('3.2 How Receiving Stock Works', 2)
add_para('Stock is only added to inventory for the amount entered in **Received Units** — not the ordered quantity. This lets a purchase be partially received (e.g. 50 of 100 ordered arrive today) without overstating stock.')
doc.add_paragraph()
add_table(
    ['Event', 'What Happens'],
    [
        ['A purchase item is created with Received Units > 0', 'Inventory.quantity increases by that amount; an "in" movement is logged, reason "Purchase received"'],
        ['Received Units is later increased/decreased on an existing item', 'Inventory.quantity changes by the difference; an "in" or "out" movement is logged, reason "Purchase receipt adjusted"'],
        ['A purchase item is deleted', 'Inventory.quantity decreases by its Received Units; an "out" movement is logged, reason "Purchase item deleted"'],
    ],
    [2.8, 4.2]
)
add_note('Every movement created from a Purchase references the Purchase it came from, so it can always be traced back to the original order and supplier.')

set_heading('3.3 Receipt & Payment Status', 2)
add_para('Two independent status badges appear on the Purchases list:')
add_table(
    ['Status', 'Meaning'],
    [
        ['Receipt: Pending', 'No items received yet'],
        ['Receipt: Partial', 'Some, but not all, ordered quantity has been received'],
        ['Receipt: Received', 'Every line item\'s Received Units meets or exceeds its Ordered Qty'],
        ['Payment: Unpaid / Partial / Paid', 'Independent of receipt — tracks what has been paid to the supplier'],
    ],
    [2.5, 4.5]
)
add_para('The **Due Date** column turns red automatically when it is in the past and the purchase is not yet fully paid, flagging overdue supplier payments at a glance.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# INVENTORY IN/OUT
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
set_heading('4. Inventory In/Out', 1)
add_para('Location: **Home → Inventory Report** (Reports section)')
add_para('Files: `InventoryReport.php`, `InventoryReportExport.php`, `InventoryMovement.php`')

set_heading('4.1 The Four Sources of Movement', 2)
add_para('Every legitimate way stock can change now writes a record to the `inventory_movements` table:')
add_table(
    ['Source', 'Type', 'Reason Logged', 'Who Can Trigger It'],
    [
        ['Purchase received / adjusted / deleted', 'in / out', 'Purchase received · Purchase receipt adjusted · Purchase item deleted', 'Anyone who can manage Purchases'],
        ['POS Sale', 'out', 'Sale', 'Cashier'],
        ['Voided Sale (manager-approved)', 'in', 'Void', 'Admin approval required'],
        ['Manual Stock edit (Edit Product)', 'in / out', 'Manual Stock Adjustment', 'super_admin, Ops Sup only'],
        ['New Product opening stock', 'in', 'Initial Stock', 'Anyone who can create Products'],
    ],
    [2.6, 0.8, 2.6, 1.7]
)

set_heading('4.2 Reading the Report', 2)
add_para('Each row shows: **Date**, **Product**, **Type** (In/Out badge), **Quantity**, **Reason**, and **Notes**. It can be filtered by type (In or Out) and by date range, and supports CSV export by preset period (Today, This Week, This Month, custom range, etc.) via the **Export to CSV** button.')
doc.add_paragraph()
add_para('Because every movement carries a `reference_id` / `reference_type`, each row can be traced back to the exact Purchase, Sale, or Void that caused it.')
doc.add_paragraph()

set_heading('4.3 End-to-End Example', 2)
add_para('A single product\'s full lifecycle, showing what each step writes to the log:')
add_table(
    ['Step', 'Action', 'Inventory Movement Logged'],
    [
        ['1', 'Product "Cement" created with 0 opening stock', 'None (0 quantity, nothing to log)'],
        ['2', 'Purchase from Supplier A: 100 pcs received', 'in · 100 · Purchase received'],
        ['3', 'Customer buys 30 pcs at POS', 'out · 30 · Sale'],
        ['4', 'That sale is voided by a manager', 'in · 30 · Void'],
        ['5', 'Admin corrects a miscount to 80 pcs (was 100)', 'out · 20 · Manual Stock Adjustment'],
        ['6', 'Current stock', '80 pcs — matches the sum of all movements above'],
    ],
    [0.5, 3.4, 3.0]
)
add_success('At any point, current stock should equal the running total of every "in" minus every "out" ever logged for that product. This is the check used to verify the report is complete.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# ROLES SUMMARY
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
set_heading('5. Roles & Permissions Summary', 1)
add_table(
    ['Role', 'Manage Products', 'Overwrite Stock Field', 'Manage Purchases', 'View Inventory Report'],
    [
        ['super_admin', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Ops Sup', 'Depends on assigned permissions', 'Yes', 'Depends on assigned permissions', 'Depends on assigned permissions'],
        ['Warehouse clerk', 'Create / edit (not Stock field)', 'No — read only', 'Depends on assigned permissions', 'Depends on assigned permissions'],
        ['cashier', 'No', 'No', 'No', 'No'],
        ['Sales Rep', 'No', 'No', 'No', 'No'],
        ['Driver', 'No', 'No', 'No', 'No'],
    ],
    [1.7, 1.7, 1.7, 1.5, 1.4]
)
add_note('Exact per-role access is managed under Roles & Permissions and can be adjusted independently of what is listed here as the current default.')
doc.add_paragraph()

set_heading('5.1 The One Rule to Remember', 2)
add_para('**Staff add stock through Purchases. Only admins/supervisors can override the stock count directly, and every change — from either path — is now logged and traceable in the Inventory In/Out Report.**')

# ════════════════════════════════════════════════════════════════════════════
doc.save('Tri-E Enterprises - Products, Purchases and Inventory In-Out Documentation.docx')
print('Saved.')
