"""
Generates user-guide.docx from the TOS User Guide content.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

HEADING_COLORS = {
    1: RGBColor(0x1E, 0x3A, 0x5F),
    2: RGBColor(0x2E, 0x6D, 0xA4),
    3: RGBColor(0x2E, 0x86, 0xAB),
    4: RGBColor(0x45, 0x8B, 0x74),
}

def set_heading(text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.color.rgb = HEADING_COLORS.get(level, RGBColor(0, 0, 0))
    run.font.bold = True
    run.font.size = Pt({1: 16, 2: 13, 3: 11.5, 4: 11}.get(level, 11))
    return h

def shade_cell(cell, hex_color='D9EAF7'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], 'D9EAF7')
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                shade_cell(cells[ci], 'F5FAFF')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run('ℹ  ' + text)
    run.font.italic = True
    run.font.size   = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x6E, 0x91)
    p.paragraph_format.left_indent = Inches(0.2)
    return p

def add_warning(text):
    p = doc.add_paragraph()
    run = p.add_run('⚠  ' + text)
    run.font.italic = True
    run.font.size   = Pt(10)
    run.font.color.rgb = RGBColor(0xB4, 0x5A, 0x09)
    p.paragraph_format.left_indent = Inches(0.2)
    return p

def add_bullet(text, level=0):
    import re
    p = doc.add_paragraph(style='List Bullet')
    text = re.sub(r'^[-*]\s*', '', text)
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p

def add_numbered(text):
    import re
    p = doc.add_paragraph(style='List Number')
    text = re.sub(r'^\d+\.\s*', '', text)
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    return p

def add_para(text):
    import re
    p = doc.add_paragraph()
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F0F0F0')
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_heading('TOS User Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
title.runs[0].font.size = Pt(28)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Tri-E Enterprises — Internal Operations System')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = desc.add_run('For Cashiers, Admin Staff, and Supervisors')
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
r2.font.italic = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. LOGGING IN
# ════════════════════════════════════════════════════════════════════════════
set_heading('1. Logging In', 1)
for step in [
    'Open your browser and go to the system URL provided by your administrator.',
    'Enter your email address and password.',
    'Click Log In.',
]:
    add_numbered(step)
add_note('If you forgot your password, contact your system administrator to reset it.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 2. NAVIGATING THE ADMIN PANEL
# ════════════════════════════════════════════════════════════════════════════
set_heading('2. Navigating the Admin Panel', 1)
add_para('After logging in you will see the admin panel. The left sidebar contains the navigation menu grouped by module:')
add_table(
    ['Group', "What's Inside"],
    [
        ['(Home)',                  'Dashboard with key metrics'],
        ['Inventory & Sales',       'Categories, Products, Suppliers, Customers, Sales, Purchases, Quotations'],
        ['Delivery Management',     'Deliveries, Drivers'],
        ['Finance',                 'Expenses, Expense Categories'],
        ['Fleet Management',        'Vehicles, Service Records, Service Requests'],
        ['Attendance Management',   'Attendance, Leave Types, Leave Requests'],
        ['Payroll',                 'Payrolls'],
        ['Reports',                 'Sales Report, Inventory Report, Aging Report, and more'],
        ['Settings',                'Users, Roles & Permissions, Audit Logs'],
    ],
    [2.2, 4.0]
)
add_note('To open the POS terminal, click Point of Sale in the sidebar or navigate to /pos.')
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 3. POINT OF SALE
# ════════════════════════════════════════════════════════════════════════════
set_heading('3. Point of Sale (POS)', 1)
add_para('The POS terminal is used for all counter sales. It runs in your browser — you do not need to reload the page between transactions.')
doc.add_paragraph()

# 3.1
set_heading('3.1 Opening the Cash Register', 2)
add_para('You must open the register before processing any sale.')
for step in [
    'When you first open the POS, a prompt appears asking for your Opening Amount.',
    'Count the physical cash in the drawer and enter the total (e.g., 500.00).',
    'Click Open Register.',
]:
    add_numbered(step)
add_para('The register is now open and tied to your user account. You are ready to process sales.')
add_note('Only one open session per user is allowed at a time.')
doc.add_paragraph()

# 3.2
set_heading('3.2 Browsing and Searching Products', 2)
for b in [
    'Product cards are displayed in the centre of the screen, sorted by stock level (highest first).',
    'Use the search bar at the top to find a product by name.',
    'Use the category filter buttons to narrow down products by category.',
    'Products with zero stock are shown with an out-of-stock indicator and cannot be added to the cart.',
]:
    add_bullet(b)
doc.add_paragraph()

# 3.3
set_heading('3.3 Adding Products to the Cart', 2)
add_para('**Standard products:**')
for step in [
    'Click the product card.',
    'If the product is sold by weight, volume, or length (kg, g, L, mL, m, ft), a modal will appear — enter the measured quantity and click Add to Cart.',
    'For all other units (pieces, bags, boxes, etc.), the item is added directly with a quantity of 1.',
]:
    add_numbered(step)
doc.add_paragraph()
add_para('**Manual / custom items** (items not in the catalogue):')
for step in [
    'Click Add Manual Item (or the custom item button in the cart panel).',
    'Enter: Item name, Unit, Unit Price, and Quantity.',
    'Click Add — the item is added to the cart but does not affect inventory.',
]:
    add_numbered(step)
doc.add_paragraph()

# 3.4
set_heading('3.4 Managing the Cart', 2)
add_para('Once items are in the cart:')
add_table(
    ['Action', 'How'],
    [
        ['Increase quantity',     'Click the + button next to the item'],
        ['Decrease quantity',     'Click the − button (minimum 0.01)'],
        ['Type a quantity',       'Click directly on the quantity number and type a new value'],
        ['Remove an item',        'Click the trash / × icon on that line'],
        ['Clear the entire cart', 'Click Clear Cart — a confirmation prompt will appear'],
    ],
    [2.2, 4.0]
)
add_note('The cart total updates automatically as you make changes.')
doc.add_paragraph()

# 3.5
set_heading('3.5 Selecting a Customer', 2)
add_para('Before charging, you can attach a customer to the sale (optional for walk-in customers).')
doc.add_paragraph()
add_para('**Search for an existing customer:**')
for step in [
    'Click the Customer field or search box.',
    'Type the customer\'s name — matching results appear as you type.',
    'Click the customer name to select them.',
]:
    add_numbered(step)
doc.add_paragraph()
add_para('**Add a new customer on the spot:**')
for step in [
    'Click Add New Customer (or the + icon next to the customer field).',
    'Fill in: Name (required), Phone, Email, Address.',
    'Click Save — the new customer is automatically selected.',
]:
    add_numbered(step)
doc.add_paragraph()
add_note('Walk-in sale: Leave the customer field blank. The sale will be recorded as a walk-in.')
doc.add_paragraph()

# 3.6
set_heading('3.6 Processing a Payment', 2)
for step in [
    'Review the cart items and total.',
    'Click the Charge button.',
    'The payment modal opens showing the total amount due.',
]:
    add_numbered(step)
doc.add_paragraph()
add_para('**Select a payment method:**')
add_table(
    ['Method', 'When to use'],
    [
        ['Cash',          'Customer pays with physical cash'],
        ['GCash',         'Customer pays via GCash QR or transfer'],
        ['Maya',          'Customer pays via Maya'],
        ['Card',          'Credit or debit card payment'],
        ['Bank Transfer', 'Direct bank transfer'],
        ['Check',         'Payment by cheque'],
    ],
    [2.0, 4.2]
)
add_para('**For cash payments:** Enter the Amount Received from the customer. The system automatically calculates the change.')
doc.add_paragraph()
for step in [
    'Click Confirm Payment.',
    'The sale is saved, inventory is updated, and a success screen appears.',
]:
    add_numbered(step)
doc.add_paragraph()

# 3.7
set_heading('3.7 Payment Terms (Credit Sales)', 2)
add_para('Payment terms allow a customer to pay later within an agreed number of days.')
doc.add_paragraph()
add_warning('Only use payment terms for trusted customers with an approved credit arrangement.')
doc.add_paragraph()
add_para('**How to apply payment terms:**')
for step in [
    'In the payment modal, toggle on Add Payment Terms.',
    'Select the credit period: 5, 10, 15, 30, or 60 days.',
    'The system automatically calculates the Due Date based on today\'s date.',
    'Select a payment method and click Confirm Payment.',
]:
    add_numbered(step)
doc.add_paragraph()
add_para('**What happens after a credit sale is saved:**')
for b in [
    'The sale is saved with payment status = unpaid.',
    'The sale is not counted in the register\'s cash totals (no money has been collected yet).',
    'The sale appears in the Sales list with a red "unpaid" badge.',
    'When the customer pays later, a supervisor marks it as paid from the admin panel (see Section 4.3).',
]:
    add_bullet(b)
add_warning('Never apply payment terms to a cash or instant payment. Payment terms are only for genuine credit arrangements.')
doc.add_paragraph()

# 3.8
set_heading('3.8 Printing a Receipt', 2)
add_para('After a successful payment, two receipt options are shown:')
add_table(
    ['Option', 'Use'],
    [
        ['Delivery Receipt', 'Customer\'s order will be delivered — includes driver/receiver signature lines'],
        ['Pick Up Receipt',  'Customer picks up at the counter — includes released-by/picked-up-by signature lines'],
    ],
    [2.0, 4.2]
)
for step in [
    'Click the receipt type button.',
    'A new tab opens with the formatted receipt.',
    'Click Print in the top-right corner.',
    'The receipt prints two copies side-by-side on A4 landscape paper: Office Copy and Customer\'s Copy, separated by a cut line.',
]:
    add_numbered(step)
doc.add_paragraph()

# 3.9
set_heading('3.9 Reprinting a Past Receipt', 2)
for step in [
    'In the POS, click the Reprint section (tab or button on the right panel).',
    'The last 50 sales are listed, most recent first.',
    'Search by receipt number or customer name.',
    'Click Print Receipt on the sale you need — choose Delivery or Pick Up format.',
]:
    add_numbered(step)
doc.add_paragraph()

# 3.10
set_heading('3.10 Creating a Quotation', 2)
add_para('A quotation saves the current cart as a price proposal without completing a sale. Inventory is not affected.')
doc.add_paragraph()
for step in [
    'Build the cart as normal (add products, set quantities).',
    'Optionally select a customer.',
    'Click Create Quotation.',
    'A modal opens — add any notes and set the validity period (default 30 days).',
    'Click Save Quotation.',
    'A print link appears immediately. Click it to open and print the quotation document.',
]:
    add_numbered(step)
doc.add_paragraph()
add_para('**Quotation statuses:**')
add_table(
    ['Status', 'Meaning'],
    [
        ['Pending',           'Waiting for admin approval'],
        ['Approved',          'Ready to convert to a sale'],
        ['Rejected',          'Declined by admin'],
        ['Converted to Sale', 'Already completed as a sale'],
        ['Expired',           'Validity period has passed'],
    ],
    [2.0, 4.0]
)

# 3.11
set_heading('3.11 Closing the Cash Register', 2)
add_para('At the end of your shift:')
for step in [
    'Click the register/cash icon or the Close Register button in the POS.',
    'Physically count all the cash in your drawer.',
    'Enter the Closing Amount (the physical cash count).',
    'Add any Notes if needed (e.g., unusual transactions, issues).',
    'Click Close Register.',
]:
    add_numbered(step)
doc.add_paragraph()
add_para('**What happens after closing:**')
for b in [
    'The system calculates the expected amount (opening cash + all cash sales collected during the session).',
    'The discrepancy is computed (positive = overage, negative = shortage).',
    'A PDF report automatically downloads in a new tab.',
]:
    add_bullet(b)
doc.add_paragraph()
add_para('**The PDF report contains:**')
for b in [
    'Session summary — Opening amount, Total sales, Cash collected, Expected cash, Closing amount, Discrepancy',
    'Full list of all sales in the session — Customer name, Time, Payment method, Payment status, Amount',
    'Session notes (if entered)',
]:
    add_bullet(b)
add_note('Save or print the PDF for your records. It can be re-downloaded anytime by your administrator.')
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 4. SALES MANAGEMENT
# ════════════════════════════════════════════════════════════════════════════
set_heading('4. Sales Management (Admin Panel)', 1)
add_para('Navigation: Inventory & Sales → Sales')
doc.add_paragraph()

# 4.1
set_heading('4.1 Viewing All Sales', 2)
add_para('The Sales list shows all recorded sales with the following columns:')
add_table(
    ['Column', 'Description'],
    [
        ['Customer', 'Customer name (or blank for walk-in)'],
        ['Date',     'Date the sale was made'],
        ['Items',    'Number of items in the sale'],
        ['Status',   'Payment status badge (green = paid, yellow = partial, red = unpaid)'],
        ['Total',    'Sale total in PHP'],
    ],
    [2.0, 4.2]
)
add_note('A grand sum of all visible sale totals is shown at the bottom of the Total column.')
doc.add_paragraph()

# 4.2
set_heading('4.2 Understanding Payment Status', 2)
add_table(
    ['Badge', 'Colour', 'Meaning'],
    [
        ['paid',    'Green',  'Payment has been fully collected'],
        ['partial', 'Yellow', 'Part of the amount has been collected'],
        ['unpaid',  'Red',    'No payment collected yet (credit/terms sale)'],
    ],
    [1.5, 1.5, 3.2]
)
add_note('Cash, GCash, Maya, Card, and Bank Transfer sales are always marked paid immediately at point of sale. Only credit/term sales start as unpaid.')
doc.add_paragraph()

# 4.3
set_heading('4.3 Marking a Credit Sale as Paid', 2)
add_para('When a customer pays their outstanding credit balance:')
for step in [
    'Go to Sales in the admin panel.',
    'Find the sale (it will have a red unpaid badge).',
    'Click the green Mark as Paid button on that row.',
    'A confirmation dialog appears showing the amount.',
    'Click Confirm.',
]:
    add_numbered(step)
doc.add_paragraph()
add_para('The sale is immediately updated:')
for b in [
    'payment_status → paid',
    'amount_paid → full total',
    'paid_date → today\'s date',
]:
    add_bullet(b)
add_warning('The Mark as Paid button only appears on credit/term sales. It will never appear on cash or instant payment sales.')
doc.add_paragraph()

# 4.4
set_heading('4.4 Downloading a Sales Summary (CSV)', 2)
add_para('The Sales Summary is a CSV file showing total sales grouped by date — useful for daily or monthly reporting.')
for step in [
    'Go to Sales in the admin panel.',
    'Click the Download Summary button (top right, with a download icon).',
    'Select your reporting period from the dialog.',
    'Click Download — the CSV file saves to your computer.',
]:
    add_numbered(step)
doc.add_paragraph()
add_para('**Available periods:**')
add_table(
    ['Option', 'Covers'],
    [
        ['Today',             'Current date only'],
        ['Yesterday',         'Previous day'],
        ['This Week',         'Monday to Sunday of the current week'],
        ['Last Week',         'Previous full week'],
        ['This Month',        'Current calendar month'],
        ['Last Month',        'Previous calendar month'],
        ['This Year',         'January to December of the current year'],
        ['Custom Date Range', 'Choose your own From and To dates'],
    ],
    [2.0, 4.0]
)
add_para('**CSV format example:**')
add_table(
    ['Date', 'Sales Count', 'Total'],
    [
        ['2026-03-01',  '8',  '24,500.00'],
        ['2026-03-02',  '5',  '15,200.00'],
        ['GRAND TOTAL', '13', '39,700.00'],
    ],
    [2.0, 1.5, 2.0]
)

# 4.5
set_heading('4.5 Editing a Sale', 2)
for step in [
    'Find the sale in the Sales list.',
    'Click the Edit (pencil) icon on that row.',
    'Make changes to the customer, date, items, or payment details.',
    'Click Save.',
]:
    add_numbered(step)
add_note('Editing a sale does not automatically adjust inventory. Contact your administrator if a product quantity needs correction.')
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 5. CUSTOMERS
# ════════════════════════════════════════════════════════════════════════════
set_heading('5. Customers', 1)
add_para('Navigation: Inventory & Sales → Customers')
doc.add_paragraph()
add_table(
    ['Field', 'Description'],
    [
        ['Name',               'Customer\'s full name or business name (required)'],
        ['Company',            'Company name (optional)'],
        ['Phone',              'Contact number'],
        ['Email',              'Email address'],
        ['Address',            'Delivery / billing address'],
        ['Payment Term Days',  '0 = cash only; e.g., 30 = Net 30 days credit'],
    ],
    [2.2, 4.0]
)
add_para('**Adding a new customer:**')
for step in [
    'Click New Customer (top right).',
    'Fill in the required Name field and any other details.',
    'Click Save.',
]:
    add_numbered(step)
add_note('Customers can also be added quickly from the POS without leaving the terminal (see Section 3.5).')
doc.add_paragraph()
add_note('If a customer has a default Payment Term Days set (e.g., 30), the system automatically calculates the due date on their sales. This can still be overridden at the POS during payment.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 6. QUOTATIONS
# ════════════════════════════════════════════════════════════════════════════
set_heading('6. Quotations', 1)
add_para('Navigation: Inventory & Sales → Quotations')
doc.add_paragraph()
add_para('Quotations are price proposals sent to customers before a sale is confirmed.')
doc.add_paragraph()
add_para('**Quotation workflow:**')
add_code('Created (Pending)  →  Approved  →  Converted to Sale\n                   →  Rejected\n                   →  Expired (if not actioned before the valid_until date)')
doc.add_paragraph()
add_para('**Approving a quotation:**')
for step in [
    'Open the quotation from the list.',
    'Review the items and total.',
    'Click Approve or Reject.',
    'Once approved, it can be converted to a sale from the POS.',
]:
    add_numbered(step)
doc.add_paragraph()
add_para('**Converting to a sale:**')
for step in [
    'Open the approved quotation.',
    'Click Convert to Sale (or use the POS link provided).',
    'The POS opens with the cart pre-filled with the quotation\'s items.',
    'Complete the payment as normal.',
]:
    add_numbered(step)
add_note('Printing: Open any quotation and click Print to generate the formatted quotation document.')
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. DELIVERIES
# ════════════════════════════════════════════════════════════════════════════
set_heading('7. Deliveries', 1)
add_para('Navigation: Delivery Management → Deliveries')
doc.add_paragraph()
add_para('After a sale is completed, a delivery record can be created to track the shipment.')
doc.add_paragraph()
add_para('**Delivery statuses:**')
add_table(
    ['Status', 'Meaning'],
    [
        ['Pending',    'Delivery created, waiting for driver assignment'],
        ['Assigned',   'Driver has been assigned'],
        ['Picked Up',  'Driver has collected the goods'],
        ['In Transit', 'Goods are on the way to the customer'],
        ['Delivered',  'Successfully delivered to the customer'],
        ['Failed',     'Delivery was unsuccessful'],
        ['Returned',   'Goods were returned'],
    ],
    [2.0, 4.2]
)
add_para('**Creating a delivery:**')
for step in [
    'Go to Deliveries and click New Delivery.',
    'Select the Sale this delivery is for.',
    'Assign a Driver.',
    'Enter the delivery address, distance (optional), and any notes.',
    'Click Save.',
]:
    add_numbered(step)
add_note('Open any delivery and click Print Receipt to generate the driver\'s delivery document with signature lines.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 8. REPORTS
# ════════════════════════════════════════════════════════════════════════════
set_heading('8. Reports', 1)
add_para('Navigation: Reports (in the left sidebar)')
doc.add_paragraph()

set_heading('8.1 Sales Report', 2)
add_para('Provides a detailed table of all sales filtered by date range.')
for step in [
    'Go to Reports → Sales Report.',
    'Use the date filter (From / Until) to set your reporting period.',
    'The table shows each sale: date, customer, item count, and total.',
    'A sum of all totals is shown at the bottom.',
]:
    add_numbered(step)
doc.add_paragraph()
add_para('**Exporting to CSV:**')
for step in [
    'Click Export to CSV (top right).',
    'Select a period or enter a custom date range.',
    'Click Export — a CSV file downloads with each individual sale as a row.',
]:
    add_numbered(step)
doc.add_paragraph()

set_heading('8.2 Inventory Report', 2)
add_para('Shows current stock levels and inventory movements.')
for step in [
    'Go to Reports → Inventory Report.',
    'Filter by product, category, or date range as needed.',
    'Export to CSV if required.',
]:
    add_numbered(step)
doc.add_paragraph()

set_heading('8.3 Aging Report', 2)
add_para('Shows outstanding unpaid balances from customers, grouped by how long they have been overdue. Use this report to identify customers who need to be followed up for payment.')
add_table(
    ['Aging Bucket', 'Overdue Days'],
    [
        ['Current',       'Not yet due'],
        ['1–30 Days',     '1 to 30 days overdue'],
        ['31–60 Days',    '31 to 60 days overdue'],
        ['61–90 Days',    '61 to 90 days overdue'],
        ['Over 90 Days',  'More than 90 days overdue'],
    ],
    [2.5, 2.5]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 9. COMMON QUESTIONS & TROUBLESHOOTING
# ════════════════════════════════════════════════════════════════════════════
set_heading('9. Common Questions & Troubleshooting', 1)

qa = [
    (
        'The POS won\'t let me process a sale — it says the register is closed.',
        'You need to open the cash register first. Click Open Register, enter your opening cash amount, and click confirm.'
    ),
    (
        'A product shows "Out of Stock" but I know we have stock.',
        'Check the product\'s inventory in Inventory & Sales → Products. The quantity may need to be updated via a purchase receipt or manual adjustment by an admin.'
    ),
    (
        'I accidentally processed a sale with the wrong amount or wrong items.',
        'Go to Sales in the admin panel, find the sale, and click Edit. Correct the details and save. For serious errors, contact your administrator.'
    ),
    (
        'The customer paid their credit balance but I can\'t find the Mark as Paid button.',
        'The button only appears on sales that have payment terms set (red "unpaid" badge). Cash or instant payment sales do not have this button because they are already paid.'
    ),
    (
        'I closed the register but the PDF report didn\'t download.',
        'Check if your browser blocked the pop-up. Allow pop-ups for this site and try closing the register again. Alternatively, ask your administrator to re-download the report.'
    ),
    (
        'The register total doesn\'t match my sales receipts.',
        'Credit/term sales are not counted in the register total until they are marked as paid. This is expected behaviour. Check the Sales list for entries with a red "unpaid" badge.'
    ),
    (
        'I can\'t find a customer when searching in the POS.',
        'Try searching with a different part of their name. If they are truly not in the system, use Add New Customer to create their profile on the spot.'
    ),
    (
        'A receipt printed incorrectly or I chose the wrong type.',
        'Use the Reprint section in the POS to reprint the receipt in the correct format (Delivery or Pick Up).'
    ),
    (
        'I need a copy of an old register closure report.',
        'Ask your administrator — they can re-download the report for any past session.'
    ),
]

for q, a in qa:
    p = doc.add_paragraph()
    run = p.add_run('Q: ' + q)
    run.font.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    p2 = doc.add_paragraph()
    run2 = p2.add_run('A: ' + a)
    run2.font.size = Pt(10.5)
    p2.paragraph_format.left_indent = Inches(0.2)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
output_path = r'c:\xampp2\htdocs\TOS\docs\user-guide.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
