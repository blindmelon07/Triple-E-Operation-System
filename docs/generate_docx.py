"""
Converts system-documentation.md to system-documentation.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Default paragraph style ───────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ── Heading colours ───────────────────────────────────────────────────────────
HEADING_COLORS = {
    1: RGBColor(0x1E, 0x3A, 0x5F),   # dark navy
    2: RGBColor(0x2E, 0x6D, 0xA4),   # medium blue
    3: RGBColor(0x2E, 0x86, 0xAB),   # steel blue
    4: RGBColor(0x45, 0x8B, 0x74),   # teal green
}

def set_heading(text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.color.rgb = HEADING_COLORS.get(level, RGBColor(0, 0, 0))
    run.font.bold = True
    run.font.size = Pt({1: 18, 2: 14, 3: 12, 4: 11}.get(level, 11))
    return h

def shade_cell(cell, hex_color='D9EAF7'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], 'D9EAF7')
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                shade_cell(cells[ci], 'F5FAFF')

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return t

def add_code_block(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F0F0F0')
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run('ℹ  ' + text)
    run.font.italic = True
    run.font.size   = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x6E, 0x91)
    p.paragraph_format.left_indent = Inches(0.2)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    # strip leading dashes/stars
    text = re.sub(r'^[-*]\s*', '', text)
    # handle inline bold **text**
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p

def add_para(text):
    p = doc.add_paragraph()
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    return p

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('TOS System Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
title.runs[0].font.size = Pt(26)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Comprehensive User & Developer Guide')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = meta.add_run('Stack: Laravel 12 · PHP 8.3 · Filament v4 · Alpine.js · MySQL')
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. POINT OF SALE
# ════════════════════════════════════════════════════════════════════════════
set_heading('1. Point of Sale (POS)', 1)
add_para('The POS is a full-featured, single-page terminal built with Alpine.js. It runs in the browser without full page reloads. Sessions can remain open all day — CSRF tokens are automatically refreshed every 90 minutes to prevent expiry errors.')
doc.add_paragraph()
add_para('**URL:** /pos   |   **Auth required:** Yes   |   **Controller:** POSController.php')
doc.add_paragraph()

# 1.1 Cash Register
set_heading('1.1 Cash Register', 2)
add_para('Before processing any sale, a cashier must open a register session.')
add_table(
    ['Action', 'Description'],
    [
        ['Open Register',  'Enter an opening cash amount. A session is created tied to the current user.'],
        ['Close Register', 'Enter the physical cash count. The system calculates the expected amount and shows any discrepancy.'],
    ],
    [2.0, 4.5]
)
add_para('The register tracks:')
for b in ['Total sales (all payment methods)', 'Total cash sales only',
          'Total number of transactions', 'Expected cash vs actual closing amount (discrepancy)']:
    add_bullet(b)
add_note('Only one open session per user at a time is allowed.')
doc.add_paragraph()

# 1.2 Adding Products to Cart
set_heading('1.2 Adding Products to Cart', 2)
add_para('**Standard products** — click a product card. If the product uses a weight/volume/length unit (kg, g, L, mL, m, ft), a modal appears to enter the measured quantity before adding.')
doc.add_paragraph()
add_para('**Manual / custom items** — for items not in the product catalogue. Enter a name, unit, unit price, and quantity. These are flagged is_manual = true and do not affect inventory.')
doc.add_paragraph()
add_para('**Stock checking:**')
for b in [
    'Products with zero stock show an out-of-stock alert (including a voice announcement via Web Speech API).',
    'Cart quantity is validated against available inventory in real time (current stock minus what is already in cart).',
]:
    add_bullet(b)
doc.add_paragraph()

# 1.3 Cart Operations
set_heading('1.3 Cart Operations', 2)
add_table(
    ['Action', 'Behaviour'],
    [
        ['+/− buttons',    'Increment or decrement quantity (minimum 0.01)'],
        ['Direct input',   'Type quantity directly; validated against available stock'],
        ['Remove item',    'Removes the line from the cart'],
        ['Clear cart',     'Requires confirmation before emptying'],
    ],
    [2.0, 4.5]
)

# 1.4 Customer Selection
set_heading('1.4 Customer Selection', 2)
for b in [
    'Search existing customers by name.',
    'Add new customer inline — name, phone, email, address. The new customer is auto-selected after saving.',
    'Leaving customer blank records the sale as a walk-in.',
]:
    add_bullet(b)
doc.add_paragraph()

# 1.5 Payment Processing
set_heading('1.5 Payment Processing', 2)
add_para('Click **Charge** to open the payment modal.')
doc.add_paragraph()
add_para('**Payment methods:**')
add_table(
    ['Method', 'Code'],
    [
        ['Cash',          'cash'],
        ['Bank Transfer', 'bank_transfer'],
        ['Check',         'check'],
        ['Credit Card',   'credit_card'],
        ['GCash',         'gcash'],
        ['Maya',          'maya'],
    ],
    [2.5, 2.5]
)
add_para('**Payment Terms (all methods):** Enable the Add Payment Terms toggle to attach a credit period to the sale. Options: 5, 10, 15, 30, or 60 days. This sets a due_date on the sale record.')
doc.add_paragraph()
add_para('**Cash change calculation:** When payment method is cash, enter the amount received — the change is calculated automatically.')
doc.add_paragraph()
add_para('When **Confirm Payment** is clicked:')
for b in [
    'A Sale record is created.',
    'SaleItem records are created for each cart line — inventory is decremented automatically for non-manual items.',
    'Cash register totals are updated.',
    'If converting from an approved quotation, the quotation status changes to converted_to_sale.',
    'A success modal appears with receipt print options (thermal or delivery receipt format).',
]:
    add_bullet(b)
doc.add_paragraph()

# 1.6 Quotations
set_heading('1.6 Quotations', 2)
add_para('From the POS, a staff member can save the cart as a quotation instead of completing a sale.')
doc.add_paragraph()
add_para('**Creating a quotation:**')
for i, b in enumerate([
    'Open the Quotation modal.',
    'Add optional notes and set validity period (default 30 days).',
    'Click Save Quotation — a Quotation record is created with the cart items. Inventory is not affected.',
    'An email notification is sent to admins with the approve_quotation permission.',
    'A print link is shown immediately after creation.',
], 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(b).font.size = Pt(10.5)
doc.add_paragraph()
add_para('**Converting a quotation to a sale:**')
for i, b in enumerate([
    'Admin approves the quotation in the admin panel (status → approved).',
    'Open POS with ?quotation_id=X in the URL — cart is pre-filled with the quotation\'s items.',
    'Complete the sale normally — quotation status changes to converted_to_sale.',
], 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(b).font.size = Pt(10.5)
doc.add_paragraph()
add_para('**Quotation statuses:**')
add_table(
    ['Status', 'Meaning'],
    [
        ['Pending',          'Awaiting admin approval'],
        ['Approved',         'Ready to convert to a sale'],
        ['Rejected',         'Declined by admin'],
        ['Converted to Sale','Sale was completed from this quotation'],
        ['Expired',          'Past the valid-until date'],
    ],
    [2.0, 4.0]
)

# 1.7 Receipt Reprinting
set_heading('1.7 Receipt Reprinting', 2)
add_para('The Reprint section shows the last 50 sales. Search by receipt number or customer name, then reprint as either a thermal receipt or delivery receipt format.')
doc.add_paragraph()

# 1.8 POS API Endpoints
set_heading('1.8 POS API Endpoints', 2)
add_table(
    ['Method', 'URL', 'Purpose'],
    [
        ['GET',  '/pos',                             'Load POS terminal'],
        ['POST', '/pos/register/open',               'Open cash register session'],
        ['POST', '/pos/register/close',              'Close cash register session'],
        ['GET',  '/pos/register/status',             'Check if session is open'],
        ['POST', '/pos/customer',                    'Create new customer inline'],
        ['POST', '/pos/quotation',                   'Create a quotation'],
        ['POST', '/pos/complete-sale',               'Complete a sale transaction'],
        ['GET',  '/pos/recent-sales',                'Fetch last 50 sales for reprint'],
        ['GET',  '/pos/print-receipt/{sale}',        'Print receipt view'],
        ['GET',  '/pos/quotation/{quotation}/print', 'Print quotation view'],
        ['GET',  '/pos/csrf-token',                  'Refresh CSRF token (long sessions)'],
    ],
    [0.8, 2.5, 3.0]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 2. INVENTORY & SALES
# ════════════════════════════════════════════════════════════════════════════
set_heading('2. Inventory & Sales', 1)

# 2.1 Categories
set_heading('2.1 Categories', 2)
add_para('Simple lookup table for organising products into groups displayed on the POS.')
add_table(
    ['Field', 'Type', 'Notes'],
    [['name', 'string', 'Required; shown on product cards in POS']],
    [1.5, 1.2, 3.5]
)

# 2.2 Products
set_heading('2.2 Products', 2)
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['name',        'string',            'Required'],
        ['category_id', 'FK → Category',     'Required'],
        ['supplier_id', 'FK → Supplier',     'Required'],
        ['price',       'decimal(10,2)',      'Selling price (₱)'],
        ['cost_price',  'decimal(10,2)',      'Purchase cost; used for profit calculations'],
        ['quantity',    'numeric',            'Tracked via linked Inventory record'],
        ['unit',        'enum (ProductUnit)', 'See unit list below'],
    ],
    [1.5, 1.8, 3.0]
)
add_para('**Product Units:**')
add_table(
    ['Unit', 'Type'],
    [
        ['Piece',                          'Unit'],
        ['Bag, Box, Bundle, Knot, Tube',   'Package'],
        ['Kilo, Gram',                     'Weight'],
        ['Liter, Milliliter',              'Liquid'],
        ['Meter, Foot',                    'Length'],
        ['Cubic Meter',                    'Volume'],
    ],
    [3.0, 2.5]
)
add_note('Weight / liquid / length units trigger the measurement modal on the POS.')
doc.add_paragraph()
add_para('**Calculated fields (read-only):**')
add_table(
    ['Field', 'Formula'],
    [
        ['profit_margin',   '(price − cost_price) / price × 100'],
        ['profit_per_unit', 'price − cost_price'],
    ],
    [2.0, 4.0]
)
add_note('If cost_price is not set, the system defaults to 70% of the selling price for margin estimates.')
doc.add_paragraph()

# 2.3 Suppliers
set_heading('2.3 Suppliers', 2)
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['name',               'string',  'Required'],
        ['contact_person',     'string',  ''],
        ['phone',              'string',  ''],
        ['email',              'string',  ''],
        ['address',            'text',    ''],
        ['payment_term_days',  'integer', '0 = COD; otherwise Net X days (auto-fills due_date on purchases)'],
    ],
    [1.8, 1.2, 3.3]
)

# 2.4 Customers
set_heading('2.4 Customers', 2)
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['name',               'string',  'Required'],
        ['contact_person',     'string',  ''],
        ['company',            'string',  ''],
        ['phone',              'string',  ''],
        ['email',              'string',  ''],
        ['address',            'text',    ''],
        ['payment_term_days',  'integer', '0 = COD; otherwise Net X days (auto-fills due_date on sales)'],
    ],
    [1.8, 1.2, 3.3]
)

# 2.5 Sales
set_heading('2.5 Sales', 2)
add_para('Sales are created through the POS. They can also be created manually in the admin panel.')
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['customer_id',            'FK → Customer', 'Nullable (walk-in)'],
        ['cash_register_session_id','FK',           'Nullable'],
        ['date',                   'date',          'Auto-set to current date'],
        ['total',                  'decimal(10,2)', 'Sum of all sale items'],
        ['payment_method',         'string',        'cash / bank_transfer / check / credit_card / gcash / maya'],
        ['payment_term_days',      'integer',       'Nullable; set if terms applied'],
        ['due_date',               'date',          'Auto-calculated from customer or POS terms'],
        ['payment_status',         'string',        'unpaid / partial / paid'],
        ['amount_paid',            'decimal(10,2)', 'Amount collected so far'],
        ['paid_date',              'date',          'Date fully paid'],
    ],
    [2.0, 1.5, 2.8]
)
add_para('**Calculated accessors:**')
add_table(
    ['Accessor', 'Description'],
    [
        ['balance',       'total − amount_paid'],
        ['days_overdue',  'Days past due_date (null if current or paid)'],
        ['aging_bucket',  'Current / 1–30 Days / 31–60 Days / 61–90 Days / Over 90 Days'],
    ],
    [2.0, 4.0]
)
add_note('Inventory is decremented automatically when a SaleItem is saved (model hook). Manual items are excluded.')
doc.add_paragraph()

# 2.6 Purchases
set_heading('2.6 Purchases', 2)
add_para('Purchases record stock received from suppliers.')
doc.add_paragraph()
add_para('**Purchase header fields:**')
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['supplier_id',     'FK → Supplier', 'Required'],
        ['date',            'date',          'Purchase date'],
        ['total',           'decimal(10,2)', 'Auto-calculated from received items only'],
        ['due_date',        'date',          'Auto-calculated from supplier payment terms'],
        ['payment_status',  'string',        'unpaid / partial / paid'],
        ['amount_paid',     'decimal(10,2)', ''],
        ['paid_date',       'date',          ''],
    ],
    [1.8, 1.3, 3.2]
)
add_para('**Purchase item fields (per line):**')
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['product_id',         'FK → Product',      'Required'],
        ['unit',               'enum (ProductUnit)', 'Auto-filled from product'],
        ['quantity',           'integer',            'Ordered quantity'],
        ['quantity_received',  'integer',            'Actually received (default 0)'],
        ['price',              'decimal(10,2)',      'Unit cost price'],
    ],
    [1.8, 1.5, 3.0]
)
add_para('**Partial Receipt Workflow:**')
for b in [
    'Set quantity = total ordered.',
    'Set quantity_received = what physically arrived.',
    'Only quantity_received drives inventory and the purchase total.',
]:
    add_bullet(b)
doc.add_paragraph()
add_para('**Receipt status (auto-calculated):**')
add_table(
    ['Status', 'Condition'],
    [
        ['Pending',  'No items received (quantity_received = 0 for all lines)'],
        ['Partial',  'At least one item received but not all'],
        ['Received', 'All items fully received (quantity_received ≥ quantity for all lines)'],
    ],
    [1.5, 4.5]
)
add_para('**Inventory impact:**')
add_table(
    ['Event', 'Effect'],
    [
        ['PurchaseItem created', 'inventory += quantity_received'],
        ['PurchaseItem updated', 'inventory adjusted by delta in quantity_received'],
        ['PurchaseItem deleted', 'inventory -= quantity_received'],
    ],
    [2.5, 3.5]
)

# 2.7 Quotations
set_heading('2.7 Quotations', 2)
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['quotation_number', 'string',       'Auto-generated: QT-YYYYMMDD-XXXX'],
        ['customer_id',      'FK → Customer','Required'],
        ['date',             'date',         'Issue date'],
        ['valid_until',      'date',         'Expiry date'],
        ['total',            'decimal(10,2)','Sum of quotation items'],
        ['notes',            'text',         'Terms, conditions, remarks'],
        ['status',           'enum',         'See quotation statuses in Section 1.6'],
        ['created_by',       'FK → User',    'Auto-set to authenticated user'],
    ],
    [1.8, 1.4, 3.1]
)
add_note('No inventory is affected when a quotation is created or approved. Inventory only changes when the quotation is converted to a sale.')
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 3. DELIVERY MANAGEMENT
# ════════════════════════════════════════════════════════════════════════════
set_heading('3. Delivery Management', 1)

# 3.1 Deliveries
set_heading('3.1 Deliveries', 2)
add_para('Deliveries are linked to a completed Sale.')
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['sale_id',           'FK → Sale',   'Required'],
        ['driver_id',         'FK → Driver', 'Assigned driver'],
        ['status',            'enum',        'See delivery statuses below'],
        ['delivery_address',  'string',      ''],
        ['distance_km',       'decimal',     ''],
        ['notes',             'text',        ''],
        ['assigned_at',       'datetime',    ''],
        ['picked_up_at',      'datetime',    ''],
        ['delivered_at',      'datetime',    ''],
        ['rating',            'integer 1–5', 'Customer satisfaction rating'],
        ['customer_feedback', 'text',        ''],
    ],
    [1.8, 1.3, 3.2]
)
add_para('**Delivery statuses:**')
add_table(
    ['Status', 'Colour'],
    [
        ['Pending',    'Gray'],
        ['Assigned',   'Blue'],
        ['Picked Up',  'Yellow'],
        ['In Transit', 'Primary'],
        ['Delivered',  'Green'],
        ['Failed',     'Red'],
        ['Returned',   'Yellow'],
    ],
    [2.0, 2.0]
)
add_note('A Print action is available on each delivery to generate a delivery receipt for the driver.')
doc.add_paragraph()

# 3.2 Drivers
set_heading('3.2 Drivers', 2)
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['name',          'string',  'Required'],
        ['phone',         'string',  ''],
        ['license_number','string',  ''],
        ['vehicle_type',  'string',  ''],
        ['vehicle_plate', 'string',  ''],
        ['is_active',     'boolean', 'Default true; inactive drivers hidden from delivery assignment'],
    ],
    [1.8, 1.2, 3.3]
)
add_para('**Calculated accessors:**')
add_table(
    ['Accessor', 'Description'],
    [
        ['delivery_count', 'Total number of deliveries'],
        ['average_rating', 'Mean rating across completed deliveries'],
        ['success_rate',   '% of deliveries with status Delivered'],
    ],
    [2.0, 4.0]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 4. FINANCE
# ════════════════════════════════════════════════════════════════════════════
set_heading('4. Finance', 1)

# 4.1 Expenses
set_heading('4.1 Expenses', 2)
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['reference_number',     'string',       'Auto-generated: EXP-YYYYMMDD-XXXX'],
        ['expense_category_id',  'FK → Category',''],
        ['expense_date',         'date',         'Cannot be a future date'],
        ['amount',               'decimal(10,2)','Required, min ₱0.01'],
        ['payment_method',       'string',       'cash / bank_transfer / check / credit_card / gcash / maya'],
        ['payee',                'string',       'Who was paid'],
        ['description',          'text',         ''],
        ['receipt_path',         'file',         'Image or PDF, max 5 MB'],
        ['status',               'string',       'pending / approved / rejected (default approved)'],
        ['user_id',              'FK → User',    'Who recorded the expense'],
    ],
    [2.0, 1.4, 2.9]
)
add_note('Only approved expenses are included in the Accounting Service total expense calculations.')
doc.add_paragraph()

# 4.2 Expense Categories
set_heading('4.2 Expense Categories', 2)
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['name',        'string',  'Required'],
        ['description', 'text',    ''],
        ['is_active',   'boolean', 'Inactive categories are hidden from the expense form'],
    ],
    [1.5, 1.2, 3.5]
)

# 4.3 Financial Overview
set_heading('4.3 Financial Overview (Accounting Service)', 2)
add_para('The Accounting Service (app/Services/AccountingService.php) powers the dashboard Financial Overview widget and can compute metrics for any date range.')
doc.add_paragraph()
add_para('**Revenue & Collections:**')
add_table(
    ['Metric', 'Description'],
    [
        ['Total Revenue',       'Sum of all sales.total in period'],
        ['Total Collections',   'Sum of sales.amount_paid in period'],
        ['Accounts Receivable', 'Revenue − Collections (outstanding balances)'],
    ],
    [2.2, 4.0]
)
add_para('**Cost & Purchasing:**')
add_table(
    ['Metric', 'Description'],
    [
        ['COGS',              'Sum of cost_price × quantity for items sold (defaults to 70% of selling price if cost_price not set)'],
        ['Total Purchases',   'Sum of all purchases.total in period'],
        ['Accounts Payable',  'Purchase total − amount paid (outstanding payables)'],
    ],
    [2.0, 4.2]
)
add_para('**Profitability:**')
add_table(
    ['Metric', 'Formula'],
    [
        ['Gross Profit',          'Revenue − COGS'],
        ['Gross Profit Margin',   'Gross Profit ÷ Revenue × 100'],
        ['Operating Costs',       'Approved Expenses + Maintenance Costs'],
        ['Operating / Net Profit','Gross Profit − Operating Costs'],
        ['Net Profit Margin',     'Net Profit ÷ Revenue × 100'],
    ],
    [2.5, 3.7]
)
add_para('**Predefined periods:** today, yesterday, this week, last week, this month, last month, this quarter, last quarter, this year, last year.')
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 5. FLEET MANAGEMENT
# ════════════════════════════════════════════════════════════════════════════
set_heading('5. Fleet Management', 1)

# 5.1 Vehicles
set_heading('5.1 Vehicles', 2)
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['plate_number',       'string',   'Unique, required'],
        ['make',               'string',   'Manufacturer'],
        ['model',              'string',   ''],
        ['year',               'integer',  '1990 to current year'],
        ['color',              'string',   ''],
        ['vin',                'string',   'Unique chassis number'],
        ['engine_number',      'string',   ''],
        ['fuel_type',          'enum',     'gasoline / diesel / electric / hybrid'],
        ['transmission',       'enum',     'automatic / manual'],
        ['current_mileage',    'integer',  'Updated after each service record'],
        ['acquisition_date',   'date',     ''],
        ['acquisition_cost',   'decimal',  ''],
        ['status',             'enum',     'active / maintenance / inactive / sold'],
        ['assigned_driver_id', 'FK → User','Optional driver assignment'],
        ['notes',              'text',     ''],
    ],
    [1.8, 1.3, 3.2]
)
add_para('**Calculated accessors:**')
add_table(
    ['Accessor', 'Description'],
    [
        ['full_name',              '{year} {make} {model}'],
        ['total_maintenance_cost', 'Sum of all service record costs for this vehicle'],
        ['last_maintenance_date',  'Date of most recent service'],
        ['maintenance_due',        'True if mileage or date exceeds the next service threshold'],
    ],
    [2.2, 4.0]
)

# 5.2 Service Types
set_heading('5.2 Service Types', 2)
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['name',                       'string',  'Required (e.g. "Oil Change", "Tire Rotation")'],
        ['description',                'text',    ''],
        ['recommended_interval_km',    'integer', 'Trigger after this many kilometres'],
        ['recommended_interval_months','integer', 'Or after this many months'],
        ['is_active',                  'boolean', ''],
    ],
    [2.2, 1.2, 2.9]
)
add_note('The interval_display accessor formats as "5,000 km or 3 months" or "As needed" if no interval is set.')
doc.add_paragraph()

# 5.3 Service Records
set_heading('5.3 Service Records', 2)
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['reference_number',      'string',              'Auto-generated: MNT-YYYYMMDD-XXXX'],
        ['vehicle_id',            'FK → Vehicle',        'Required'],
        ['maintenance_type_id',   'FK → ServiceType',    'Required'],
        ['user_id',               'FK → User',           'Technician / person in charge'],
        ['maintenance_date',      'date',                'Required'],
        ['mileage_at_service',    'integer',             'Odometer reading at time of service'],
        ['parts_cost',            'decimal',             ''],
        ['labor_cost',            'decimal',             ''],
        ['service_provider',      'string',              'External shop name if applicable'],
        ['description',           'text',                'Work performed'],
        ['parts_replaced',        'text',                'List of replaced parts'],
        ['next_service_date',     'date',                ''],
        ['next_service_mileage',  'integer',             ''],
        ['status',                'enum',                'completed / pending / in_progress'],
        ['invoice_path',          'file',                'Upload of service invoice'],
    ],
    [2.0, 1.4, 2.9]
)
add_para('**Calculated accessors:** total_cost = parts_cost + labor_cost  |  is_overdue = past next_service_date or mileage')
doc.add_paragraph()

# 5.4 Service Requests
set_heading('5.4 Service Requests', 2)
add_para('Staff submit service requests which go through an approval workflow before becoming a Service Record.')
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['request_number',       'string',            'Auto-generated: REQ-YYYYMMDD-XXXX'],
        ['vehicle_id',           'FK → Vehicle',      'Required'],
        ['maintenance_type_id',  'FK → ServiceType',  ''],
        ['priority',             'enum',              'high / normal / low'],
        ['current_mileage',      'integer',           'Current odometer reading'],
        ['description',          'text',              'Description of the issue'],
        ['preferred_date',       'date',              'Requested service date'],
        ['status',               'enum',              'pending / approved / rejected / completed'],
        ['estimated_cost',       'decimal',           'Filled in during approval'],
        ['rejection_reason',     'text',              'Required if rejected'],
        ['approved_by',          'FK → User',         ''],
        ['maintenance_record_id','FK → ServiceRecord','Linked after completion'],
    ],
    [2.0, 1.5, 2.8]
)
add_note('A badge on the navigation item shows the count of pending service requests.')
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 6. ATTENDANCE MANAGEMENT
# ════════════════════════════════════════════════════════════════════════════
set_heading('6. Attendance Management', 1)

# 6.1 Attendance
set_heading('6.1 Attendance', 2)
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['date',        'date',    'Cannot be a future date'],
        ['user_id',     'FK → User','Employee (super_admin excluded)'],
        ['time_in',     'time',    ''],
        ['time_out',    'time',    ''],
        ['total_hours', 'decimal', 'Auto-calculated from time_in / time_out'],
        ['status',      'enum',    'Present / Absent / Late / Half Day / On Leave'],
        ['remarks',     'text',    ''],
        ['recorded_by', 'FK → User','Who logged this record'],
    ],
    [1.5, 1.5, 3.3]
)
add_para('**Attendance statuses:**')
add_table(
    ['Status',   'Colour'],
    [
        ['Present',   'Green'],
        ['Absent',    'Red'],
        ['Late',      'Yellow'],
        ['Half Day',  'Blue'],
        ['On Leave',  'Gray'],
    ],
    [2.0, 2.0]
)

# 6.2 Leave Types
set_heading('6.2 Leave Types', 2)
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['name',               'string',  'Required (e.g. "Sick Leave", "Vacation Leave")'],
        ['description',        'text',    ''],
        ['max_days_per_year',  'integer', 'Annual entitlement'],
        ['is_paid',            'boolean', 'Whether this leave type is compensated'],
        ['is_active',          'boolean', ''],
    ],
    [2.0, 1.2, 3.1]
)

# 6.3 Leave Requests
set_heading('6.3 Leave Requests', 2)
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['request_number',    'string',         'Auto-generated: LR-YYYYMMDD-XXXX'],
        ['user_id',           'FK → User',      'Employee making the request'],
        ['leave_type_id',     'FK → LeaveType', ''],
        ['start_date',        'date',           ''],
        ['end_date',          'date',           ''],
        ['total_days',        'integer',        'Auto-calculated'],
        ['reason',            'text',           ''],
        ['status',            'enum',           'pending / approved / rejected / cancelled'],
        ['rejection_reason',  'text',           'Required if rejected'],
        ['approved_by',       'FK → User',      ''],
    ],
    [1.8, 1.4, 3.1]
)
add_note('A badge on the navigation item shows the count of pending leave requests.')
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. PAYROLL
# ════════════════════════════════════════════════════════════════════════════
set_heading('7. Payroll', 1)
add_para('Payrolls follow a Draft → Approved → Paid workflow.')
doc.add_paragraph()
add_para('**Payroll header:**')
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['payroll_number',    'string',          'Auto-generated: PAY-YYYYMMDD-XXXX'],
        ['pay_period_start',  'date',            ''],
        ['pay_period_end',    'date',            ''],
        ['pay_period_type',   'enum',            'Daily / Weekly / Semi-Monthly'],
        ['total_gross',       'decimal',         'Auto-calculated from payroll items'],
        ['total_deductions',  'decimal',         'Auto-calculated from payroll items'],
        ['total_net',         'decimal',         'gross − deductions'],
        ['status',            'enum',            'Draft / Approved / Paid / Cancelled'],
        ['approved_by',       'FK → User',       ''],
        ['paid_at',           'datetime',        ''],
        ['notes',             'text',            ''],
    ],
    [1.8, 1.3, 3.2]
)
add_para('**Payroll items** are managed via a Relation Manager inside the Payroll view. Each item covers one employee\'s gross pay, deductions, and net pay for the period.')
doc.add_paragraph()
add_para('**Workflow rules:**')
add_table(
    ['Method', 'Condition'],
    [
        ['canBeApproved()', 'Only when status is Draft'],
        ['canBePaid()',      'Only when status is Approved'],
        ['canBeCancelled()','Cannot cancel after Paid'],
    ],
    [2.5, 3.5]
)
add_note('A badge on the navigation item shows the count of Draft payrolls awaiting approval.')
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 8. AUTHENTICATION & SECURITY
# ════════════════════════════════════════════════════════════════════════════
set_heading('8. Authentication & Security', 1)

# 8.1 Users
set_heading('8.1 Users', 2)
add_table(
    ['Field', 'Type', 'Notes'],
    [
        ['name',              'string',   'Required'],
        ['email',             'string',   'Required, unique'],
        ['password',          'string',   'Hashed'],
        ['email_verified_at', 'datetime', ''],
    ],
    [2.0, 1.2, 3.1]
)
add_para('Users are assigned one or more roles via Spatie Permissions. Users also serve as employees and are linked to Attendance, LeaveRequest, EmployeeCompensation, and PayrollItem records.')
add_note('The audit trail excludes sensitive fields: password, remember_token, email_verified_at.')
doc.add_paragraph()

# 8.2 Roles & Permissions
set_heading('8.2 Roles & Permissions', 2)
add_para('Roles are managed in the admin panel using FilamentShield (Spatie Permissions). Each role can be granted or denied access to individual resources (view, create, update, delete) and custom permissions.')
doc.add_paragraph()
add_para('**Notable custom permission:**')
add_table(
    ['Permission', 'Effect'],
    [
        ['approve_quotation', 'Allows approving or rejecting quotations and receiving email notifications for new quotations.'],
    ],
    [2.0, 4.2]
)

# 8.3 Audit Logs
set_heading('8.3 Audit Logs', 2)
add_para('All significant actions are recorded automatically via the Auditable trait used across most models. The log is read-only — no create, edit, or delete is available from the interface.')
doc.add_paragraph()
add_table(
    ['Field', 'Description'],
    [
        ['user_name',       'Name of the user who performed the action'],
        ['action',          'e.g. completed_sale, approved, opened_register'],
        ['auditable_label', 'Human-readable name of the affected record'],
        ['auditable_type',  'Model class (e.g. App\\Models\\Sale)'],
        ['old_values',      'JSON of field values before the change'],
        ['new_values',      'JSON of field values after the change'],
        ['ip_address',      'Client IP address'],
        ['user_agent',      'Browser / device string'],
        ['created_at',      'Timestamp'],
    ],
    [2.0, 4.2]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# APPENDIX A — Reference Numbers
# ════════════════════════════════════════════════════════════════════════════
set_heading('Appendix A — Reference Numbers', 1)
add_para('All auto-generated reference numbers follow this format. The sequence resets daily.')
add_table(
    ['Record', 'Format', 'Example'],
    [
        ['Quotation',       'QT-YYYYMMDD-XXXX',  'QT-20260220-0001'],
        ['Expense',         'EXP-YYYYMMDD-XXXX', 'EXP-20260220-0001'],
        ['Service Record',  'MNT-YYYYMMDD-XXXX', 'MNT-20260220-0001'],
        ['Service Request', 'REQ-YYYYMMDD-XXXX', 'REQ-20260220-0001'],
        ['Leave Request',   'LR-YYYYMMDD-XXXX',  'LR-20260220-0001'],
        ['Payroll',         'PAY-YYYYMMDD-XXXX', 'PAY-20260220-0001'],
    ],
    [2.0, 2.5, 2.0]
)

# ════════════════════════════════════════════════════════════════════════════
# APPENDIX B — Inventory Flow
# ════════════════════════════════════════════════════════════════════════════
set_heading('Appendix B — Inventory Flow Summary', 1)
add_code_block(
    "Purchase Created\n"
    "    └─ PurchaseItem saved\n"
    "           └─ quantity_received > 0  →  Inventory.quantity  +=  quantity_received\n"
    "\n"
    "Sale Completed (POS or Admin)\n"
    "    └─ SaleItem saved\n"
    "           └─ is_manual = false  →  Inventory.quantity  -=  quantity\n"
    "\n"
    "Quotation Created\n"
    "    └─ QuotationItem saved\n"
    "           └─ No inventory change (quotation is not a commitment)\n"
    "\n"
    "Quotation → Sale (Convert)\n"
    "    └─ SaleItem saved (same as Sale Completed above)"
)

# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════
output_path = r'c:\xampp2\htdocs\TOS\docs\system-documentation.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
