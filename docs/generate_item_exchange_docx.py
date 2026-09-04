"""
Generates: Tri-E Enterprises - POS Item Exchange User Guide.docx

A step-by-step guide to exchanging a single line item on an already-completed
POS sale: what the cashier does, what the manager approves, and exactly what
the system does to stock, the sale total, and the cash drawer afterwards.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

ASSETS = os.path.join(os.path.dirname(os.path.abspath(__file__)), '_exchange_assets')

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Default paragraph style ───────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ── Heading colours ───────────────────────────────────────────────────────────
HEADING_COLORS = {
    1: RGBColor(0x1E, 0x3A, 0x5F),   # dark navy
    2: RGBColor(0x2E, 0x6D, 0xA4),   # medium blue
    3: RGBColor(0x2E, 0x86, 0xAB),   # steel blue
    4: RGBColor(0x45, 0x8B, 0x74),   # teal green
}


def set_heading(text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.color.rgb = HEADING_COLORS.get(level, RGBColor(0, 0, 0))
    run.font.bold = True
    run.font.size = Pt({1: 18, 2: 14, 3: 12, 4: 11}.get(level, 11))
    return h


def shade_cell(cell, hex_color='D9EAF7'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], 'D9EAF7')
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                shade_cell(cells[ci], 'F5FAFF')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return t


def add_code_block(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F0F0F0')
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_note(text, color=RGBColor(0x44, 0x6E, 0x91), icon='ℹ'):
    p = doc.add_paragraph()
    run = p.add_run(icon + '  ' + text)
    run.font.italic = True
    run.font.size   = Pt(10)
    run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.2)
    return p


def add_warning(text):
    return add_note(text, color=RGBColor(0xA6, 0x3D, 0x2E), icon='⚠')


def add_success(text):
    return add_note(text, color=RGBColor(0x2E, 0x7D, 0x4F), icon='✔')


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    text = re.sub(r'^[-*]\s*', '', text)
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p


def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    return p


def add_para(text):
    p = doc.add_paragraph()
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    return p


def add_image(filename, width_in=5.5, caption=None):
    path = os.path.join(ASSETS, filename)
    if not os.path.exists(path):
        add_warning(f'Screenshot not found: {filename}')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.font.size = Pt(8.5)
        cr.font.italic = True
        cr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        cap.paragraph_format.space_after = Pt(10)


def add_step(number, title, body_lines):
    """A numbered step with a bold heading line and indented detail lines."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f'Step {number} — {title}')
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    for line in body_lines:
        d = doc.add_paragraph()
        d.paragraph_format.left_indent = Inches(0.3)
        d.paragraph_format.space_after = Pt(2)
        parts = re.split(r'\*\*(.*?)\*\*', line)
        for i, part in enumerate(parts):
            run = d.add_run(part)
            run.bold = (i % 2 == 1)
            run.font.size = Pt(10.5)
    return p


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('POS Item Exchange', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
title.runs[0].font.size = Pt(26)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Step-by-Step User Guide')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = meta.add_run('Tri-E Enterprises · TOS System · Point of Sale Module')
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = meta2.add_run('Prepared: September 4, 2026')
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph()
doc.add_paragraph()

intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
ri = intro.add_run(
    'Swapping one item on a sale that has already been completed —\n'
    'returning the old item to stock, issuing the new one, and settling\n'
    'the price difference with the customer.'
)
ri.font.size = Pt(11)
ri.font.italic = True
ri.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
set_heading('1. Overview', 1)

add_para(
    'A customer sometimes changes their mind after the receipt has already printed — '
    'they took the wrong size, the wrong colour, or simply want something else instead. '
    'Before this feature the only way to handle that was to void the whole sale and ring '
    'it up again from scratch.'
)
add_para(
    '**Item Exchange** lets you swap a single line item on a completed sale for a different '
    'product, without disturbing the rest of the sale. The old item goes back into stock, the '
    'new item comes out of stock, the sale total is recalculated, and the price difference is '
    'either collected from the customer or refunded to them.'
)

set_heading('Who does what', 2)
add_table(
    ['Role', 'What they can do'],
    [
        ['Cashier', 'Requests the exchange from the POS screen. Nothing changes until it is approved.'],
        ['Admin / Super Admin', 'Reviews the request, sees the price difference, then approves or rejects it.'],
    ],
    col_widths=[1.5, 4.8],
)

add_note(
    'An exchange follows the same manager-approval rule as a void. A cashier can never rewrite '
    'a finished sale on their own.'
)

set_heading('When to use which action', 2)
add_table(
    ['Situation', 'Use this'],
    [
        ['Customer wants a different product instead', 'Item Exchange'],
        ['Customer is returning one item and taking nothing in its place', 'Void Item'],
        ['The whole transaction was a mistake', 'Void Transaction'],
        ['Customer is only changing the quantity of the same product', 'Item Exchange (pick the same product, new quantity)'],
    ],
    col_widths=[3.6, 2.7],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 2. BEFORE YOU START
# ════════════════════════════════════════════════════════════════════════════
set_heading('2. Before You Start', 1)

add_para('An exchange can only be requested when all of the following are true:')

add_table(
    ['Requirement', 'Why'],
    [
        ['You have an open register session',
         'The money movement has to land in a real, open drawer.'],
        ['The sale belongs to your own open session',
         'Keeps the refund or extra payment in the same drawer that took the original money.'],
        ['The sale has not been voided',
         'There is nothing left to exchange on a voided sale.'],
        ['The line item has not been voided',
         'That item is already off the sale.'],
        ['No other request is pending on that item',
         'One item can only have one request in flight at a time.'],
        ['The replacement product has enough stock',
         'You cannot issue something you do not have.'],
    ],
    col_widths=[2.6, 3.7],
)

add_note(
    'Unlike Void Item, an exchange IS allowed on the last remaining item of a sale — the sale '
    'still has an item afterwards, so it never ends up empty.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 3. CASHIER STEPS
# ════════════════════════════════════════════════════════════════════════════
set_heading('3. Part 1 — Cashier: Requesting an Exchange', 1)

add_para(
    'These are the steps at the counter, with the customer in front of you.'
)

add_step(1, 'Open the Recent Sales list', [
    'From the POS screen, open **Recent Sales** (the same list you use for reprinting a receipt).',
    'Find the customer’s receipt. You can search by receipt number or customer name.',
])

add_step(2, 'Find the item to swap', [
    'Under each sale you will see its line items listed, one per row.',
    'Only sales from **your own open register session** show these per-item buttons.',
])

add_step(3, 'Click the blue exchange icon', [
    'Each item row has two small icons on the right:',
    '     🔵  **Blue swap arrows** — Exchange this item for another product',
    '     🔴  **Red X** — Void this item (remove it entirely)',
    'Click the **blue swap arrows** on the item the customer is returning.',
])

add_image('01-recent-sales-item-actions.png', 6.0,
          'A sale in Recent Sales. Each item row carries the blue exchange icon and the red void icon.')

add_step(4, 'Choose the replacement product', [
    'The **Exchange Item** window opens. The heading shows what is being replaced, '
    'for example: *Replacing: Hammer × 1 — ₱100.00*',
    'Type in the search box to find the new product.',
    'Each result shows its price and its current stock on hand.',
    'Click the product to select it.',
])

add_image('02-exchange-search.png', 5.2,
          'Searching for the replacement. Each result shows its price and stock on hand.')

add_step(5, 'Set quantity, unit and price', [
    '**Quantity** — how many of the new item the customer is taking.',
    '**Unit** — the selling unit (piece, meter, roll, and so on). Changing the unit '
    'automatically fills in that unit’s price.',
    '**Unit price** — filled in for you, but you can override it if needed.',
])

add_step(6, 'Check the price difference', [
    'The summary box updates as you type and shows three lines:',
    '     **Item removed** — the value coming off the sale',
    '     **Item added** — the value going onto the sale',
    '     The bold line tells you what to do with the money:',
    '          *Customer pays ₱X* — the new item costs more',
    '          *Refund customer ₱X* — the new item costs less',
    '          *No difference* — a straight, equal-value swap',
])

add_image('03-exchange-price-difference.png', 4.5,
          'Replacing ₱330.00 of 10mm bar with ₱460.00 of 12mm bar — the customer pays ₱130.00.')

add_step(7, 'Enter the reason and submit', [
    'Type a short **reason for exchange** — for example, “Customer wanted a different size”. '
    'This is required, and it is stored on the record permanently.',
    'Click **Request Exchange**.',
])

add_step(8, 'Wait for the manager', [
    'A waiting screen appears while the request sits with a manager.',
    'Admins receive a notification straight away, and the pending badge on their screen goes up by one.',
    'If you clicked by mistake, you can **Cancel** the request while it is still waiting.',
])

add_image('04-waiting-for-approval.png', 3.2,
          'The cashier’s screen while the request sits with a manager.')

add_warning(
    'Do not hand over the new item, and do not take or give any money, until the manager has '
    'approved the request. Nothing is final until then.'
)

set_heading('What you see when it is decided', 2)
add_table(
    ['Outcome', 'What happens on your screen'],
    [
        ['Approved',
         'The waiting screen closes and a message tells you exactly how much to collect or refund. '
         'The sales list refreshes with the new item and the new total.'],
        ['Rejected',
         'The waiting screen closes and the manager’s reason for rejecting is shown. '
         'The sale is left exactly as it was.'],
    ],
    col_widths=[1.3, 5.0],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 4. MANAGER STEPS
# ════════════════════════════════════════════════════════════════════════════
set_heading('4. Part 2 — Manager: Reviewing and Approving', 1)

add_para(
    'Exchange requests appear in the same approvals panel as void requests, so there is only '
    'one queue to watch.'
)

add_step(1, 'Open the approvals panel', [
    'The pending-requests button on the POS screen shows a count whenever something is waiting.',
    'Click it to open the panel.',
])

add_step(2, 'Identify the request type', [
    'Each request is tagged so you can tell them apart at a glance:',
    '     **ITEM EXCHANGE** (blue tag) — a swap',
    '     **ITEM VOID** (orange tag) — a single item being removed',
    '     No tag — the whole transaction is being voided',
])

add_step(3, 'Review the details', [
    'An exchange request shows you:',
    '     **Receipt number** and **customer**',
    '     **Remove:** the item going off the sale, with its quantity and value',
    '     **Add:** the replacement, with its quantity, unit and value',
    '     A coloured line stating **Customer pays ₱X**, **Refund ₱X**, or **No price difference**',
    '     **Sale total** as it stands right now',
    '     Who requested it, and their stated reason',
])

add_image('05-manager-approval-card.png', 6.0,
          'An exchange request as the manager sees it, with the blue ITEM EXCHANGE tag '
          'and the amount to collect.')

add_step(4, 'Approve or reject', [
    '**Approve** — the swap is applied immediately and in full (see Part 3).',
    '**Reject** — you must type a reason, which is shown back to the cashier.',
])

add_step(5, 'Settle the money at the counter', [
    'The approval message repeats the amount to collect or refund.',
    'Hand over or take the difference in cash at the counter. The system has already '
    'adjusted the drawer figures to match.',
])

add_note(
    'Only Admin and Super Admin roles can approve. A cashier who tries is refused, even if '
    'they request it themselves.'
)

set_heading('The result', 2)
add_para(
    'Once approved, the sale is updated straight away. The old item is gone from the list, the '
    'replacement is in its place, and the receipt total reflects the new figure — while the '
    'receipt number stays the same.'
)
add_image('07-sale-after-exchange.png', 6.0,
          'The same receipt after approval: 12mm bar in place of the 10mm, total now ₱560.00.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 5. WHAT THE SYSTEM DOES
# ════════════════════════════════════════════════════════════════════════════
set_heading('5. Part 3 — What the System Does on Approval', 1)

add_para(
    'Everything below happens as a **single, all-or-nothing operation**. If any part of it fails — '
    'most commonly the replacement selling out between the request and the approval — then nothing '
    'at all is applied and the request stays pending.'
)

set_heading('Stock', 2)
add_bullet('The old item’s quantity is **returned** to inventory.')
add_bullet('The new item’s quantity is **taken** from inventory.')
add_bullet('Both movements are written to the inventory movement log with the reason **“Item Exchange”**, so the stock trail stays complete.')
add_bullet('Quantities are converted to the product’s base unit first, so exchanging 1 roll for 5 meters adjusts stock correctly.')

set_heading('The sale', 2)
add_bullet('The old line item is marked as voided, with the reason recorded as **“Exchanged: …”** — it stays on the record for audit purposes rather than being deleted.')
add_bullet('The replacement is added as a **new line item** on the same sale.')
add_bullet('The sale total is recalculated: **old total − item removed + item added**.')
add_bullet('The payment status (paid / partial / unpaid) is recalculated against the new total.')
add_bullet('The receipt number does not change, and the sale still counts as one transaction.')

set_heading('The money', 2)
add_para('What happens to the drawer depends on which way the price went, and on whether the sale had already been paid in full:')

add_table(
    ['Situation', 'What the system does'],
    [
        ['New item costs more, sale was fully paid',
         'The difference is collected now and added to the register session total. The sale stays “paid”.'],
        ['New item costs more, sale still has a balance',
         'No money is taken now — the extra simply increases what the customer still owes.'],
        ['New item costs less',
         'The overpaid amount is reversed out of the session(s) that originally took it, oldest money first.'],
        ['Same price',
         'The drawer is not touched at all.'],
    ],
    col_widths=[2.5, 3.8],
)

add_note(
    'The transaction count on the register session never changes — an exchange adjusts an '
    'existing sale, it does not create a new one. Your end-of-day transaction count stays honest.'
)

set_heading('Audit trail', 2)
add_para('Every approved exchange writes an audit log entry recording:')
add_bullet('Who approved it, and when')
add_bullet('The reason given by the cashier')
add_bullet('The item removed and the item added, with both values')
add_bullet('The old total and the new total')
add_bullet('How much was collected or refunded')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 6. WORKED EXAMPLES
# ════════════════════════════════════════════════════════════════════════════
set_heading('6. Worked Examples', 1)

add_para(
    'All four examples use the same starting sale — the one shown in the screenshots throughout '
    'this guide: a walk-in customer who paid **₱430.00** cash for 2 × RSB 10MMX6M G33 (₱330.00) '
    'and 1 kilo of Concrete Nail 3" (₱100.00).'
)

set_heading('Example A — Replacement costs more', 3)
add_para(
    'The customer comes back wanting the thicker 12mm bar instead of the 10mm. '
    'This is the exchange captured in the screenshots above.'
)
add_code_block(
    'Sale total before          ₱430.00\n'
    'Remove  RSB 10MMX6M × 2    −₱330.00\n'
    'Add     RSB 12MMX6M × 2    +₱460.00\n'
    '───────────────────────────────────\n'
    'Sale total after           ₱560.00\n'
    'COLLECT FROM CUSTOMER      ₱130.00'
)
add_bullet('10mm stock goes **up** by 2; 12mm stock goes **down** by 2.')
add_bullet('Register session sales rise by ₱130; transaction count stays at 1.')
add_bullet('Sale remains **paid**, with amount paid now ₱560.')

set_heading('Example B — Replacement costs less', 3)
add_para('Same ₱430 sale, but this time they want the thinner 9mm bar (₱125.00 each) instead.')
add_code_block(
    'Sale total before          ₱430.00\n'
    'Remove  RSB 10MMX6M × 2    −₱330.00\n'
    'Add     RSB 9MM     × 2    +₱250.00\n'
    '───────────────────────────────────\n'
    'Sale total after           ₱350.00\n'
    'REFUND TO CUSTOMER          ₱80.00'
)
add_bullet('Register session sales drop by ₱80; transaction count stays at 1.')
add_bullet('Sale remains **paid**, with amount paid now ₱350.')

set_heading('Example C — Equal swap', 3)
add_para('The customer swaps the ₱100.00 of concrete nails for a different ₱100.00 product.')
add_code_block(
    'Sale total before          ₱430.00\n'
    'Sale total after           ₱430.00\n'
    'NO MONEY CHANGES HANDS'
)
add_bullet('Only stock moves: one product back in, another out.')
add_bullet('The drawer and the sale total are untouched.')

set_heading('Example D — Quantity change on the same product', 3)
add_para(
    'The customer wants 3 of the 10mm bars rather than 2. Select **the same product** as the '
    'replacement and set the quantity to 3.'
)
add_code_block(
    'Remove  RSB 10MMX6M × 2    −₱330.00\n'
    'Add     RSB 10MMX6M × 3    +₱495.00\n'
    '───────────────────────────────────\n'
    'Sale total after           ₱595.00\n'
    'COLLECT FROM CUSTOMER      ₱165.00'
)
add_note(
    'The system returns the 2 units to stock before checking whether 3 are available, so a '
    'quantity increase works correctly even when stock is tight.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. RULES AND RESTRICTIONS
# ════════════════════════════════════════════════════════════════════════════
set_heading('7. Rules and Restrictions', 1)

add_bullet('**Manager approval is always required.** A cashier can request but never apply an exchange.')
add_bullet('**Same-session only.** You can only exchange items on sales taken in your own currently-open register session. Yesterday’s sale cannot be exchanged today.')
add_bullet('**One request per item.** If a void or exchange is already pending on an item, a second request is not created.')
add_bullet('**Replacements must be catalogue products.** You cannot exchange into a manually-typed custom item.')
add_bullet('**Discounts do not carry over.** The replacement is priced as unit price × quantity. If the customer is owed a discount on the new item, lower the unit price by hand before submitting.')
add_bullet('**Stock is checked twice** — once when the cashier submits, and again under lock at the moment of approval.')
add_bullet('**Nothing is deleted.** The old line item stays on the sale, marked as exchanged, so the history is always reconstructable.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 8. TROUBLESHOOTING
# ════════════════════════════════════════════════════════════════════════════
set_heading('8. Messages and What They Mean', 1)

add_table(
    ['Message', 'Meaning / What to do'],
    [
        ['No open register session.',
         'Open your register session first, then try again.'],
        ['This item does not belong to a sale in the current register session.',
         'The sale was taken in a different session, or by a different cashier. It cannot be exchanged from here.'],
        ['This sale has already been voided.',
         'There is nothing left to exchange.'],
        ['This item has already been voided.',
         'That line is already off the sale. Nothing to do.'],
        ['A request for this item is already pending.',
         'Someone already submitted a void or exchange for this item. Wait for the manager to decide.'],
        ['Not enough stock for replacement product: …',
         'The replacement does not have enough on hand. Reduce the quantity or pick a different product.'],
        ['Replacement product no longer exists.',
         'The product was deleted between the request and the approval. Reject the request and start again.'],
        ['Request is no longer pending.',
         'Another manager already approved or rejected it.'],
        ['Unauthorized.',
         'Only Admin and Super Admin can approve or reject.'],
    ],
    col_widths=[2.7, 3.6],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 9. TECHNICAL REFERENCE
# ════════════════════════════════════════════════════════════════════════════
set_heading('9. Technical Reference', 1)
add_para('For developers and system administrators maintaining the module.')

set_heading('Routes', 2)
add_code_block(
    'POST  /pos/exchange-request-item/{saleItem}     POSController@requestItemExchange\n'
    'POST  /pos/void-requests/{voidRequest}/approve  VoidRequestController@approve\n'
    'POST  /pos/void-requests/{voidRequest}/reject   VoidRequestController@reject'
)

set_heading('Database', 2)
add_para('Exchanges reuse the existing **void_requests** table so both flows share one approval queue. Added columns:')
add_table(
    ['Column', 'Type', 'Purpose'],
    [
        ['type', 'string, default "void"', '"void" or "exchange" — which flow this request is'],
        ['replacement_product_id', 'FK products, nullable', 'The product being swapped in'],
        ['replacement_quantity', 'decimal(15,2), nullable', 'How many of it'],
        ['replacement_unit', 'string, nullable', 'Selling unit for the replacement'],
        ['replacement_unit_price', 'decimal(15,2), nullable', 'Price per unit at time of request'],
    ],
    col_widths=[1.9, 1.7, 2.7],
)

add_note(
    'VoidRequest::isItemVoid() deliberately excludes exchanges, so an exchange is never routed '
    'into the item-void approval path.'
)

set_heading('Key code paths', 2)
add_table(
    ['File', 'Responsibility'],
    [
        ['app/Http/Controllers/POSController.php',
         'requestItemExchange() — validation, session and stock guards, creates the pending request, notifies managers'],
        ['app/Http/Controllers/VoidRequestController.php',
         'approveItemExchange() — the whole swap inside one locked transaction; refundOldestFirst() is shared with the item-void flow'],
        ['app/Models/VoidRequest.php',
         'isItemExchange(), replacementProduct()'],
        ['app/Models/CashRegisterSession.php',
         'addAmount() — collects extra without incrementing the transaction count'],
        ['app/Models/SaleItem.php',
         'created() hook decrements stock automatically when the replacement line is written'],
        ['resources/views/pos/index.blade.php',
         'Exchange modal, per-item buttons, Alpine state, manager panel display'],
    ],
    col_widths=[2.6, 3.7],
)

set_heading('Accounting note', 2)
add_para(
    'When extra money is collected on an already-paid sale, the amount is added to the register '
    'session but **no SalePayment row is written**. This is deliberate and mirrors completeSale(): '
    'a fully-paid sale is already counted through the sales list, so a payment row would '
    'double-count the same money in the daily and period reports.'
)

set_heading('Tests', 2)
add_para('Covered by **tests/Feature/ItemExchangeTest.php** — 13 tests:')
add_bullet('Request validation: submission, last-item allowance, insufficient stock, wrong session, already-voided item')
add_bullet('Approval: item swap, total recalculation, stock in and out, extra collected, refund issued, price-neutral swap')
add_bullet('Rollback: a stock-out between request and approval leaves nothing half-applied')
add_bullet('Authorisation: non-admin cannot approve')
add_code_block('php artisan test tests/Feature/ItemExchangeTest.php')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 10. QUICK REFERENCE
# ════════════════════════════════════════════════════════════════════════════
set_heading('10. Quick Reference Card', 1)

set_heading('Cashier', 3)
add_numbered('Recent Sales → find the receipt')
add_numbered('Click the blue swap icon on the item')
add_numbered('Search and select the replacement product')
add_numbered('Set quantity, unit and price')
add_numbered('Read the price difference in the summary box')
add_numbered('Type the reason → Request Exchange')
add_numbered('Wait for approval — then settle the difference')

set_heading('Manager', 3)
add_numbered('Open the pending approvals panel')
add_numbered('Look for the blue ITEM EXCHANGE tag')
add_numbered('Check Remove / Add lines and the price difference')
add_numbered('Approve, or Reject with a reason')

doc.add_paragraph()
add_success('Approved exchanges take effect immediately — stock, sale total and drawer are all updated together.')
add_warning('Never hand over goods or money before the manager has approved.')

doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run('Tri-E Enterprises · TOS System · POS Item Exchange · September 2026')
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ════════════════════════════════════════════════════════════════════════════
out = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    'Tri-E Enterprises - POS Item Exchange User Guide.docx'
)
doc.save(out)
print('Saved:', out)
