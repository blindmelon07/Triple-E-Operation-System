"""
Generates TOS-Complete-User-Guide.docx — a full, screenshot-illustrated user
guide covering every module of the Tri-E Trading Operating System (TOS).

Screenshots are read from c:\\tmp\\tos_shots (captured live from the running
app via Playwright). Run generate order: capture screenshots first, then
run this script.
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

SHOTS = Path(r"c:\tmp\tos_shots")
OUTPUT = Path(r"c:\xampp2\htdocs\TOS\docs\Tri-E Enterprises - Complete User Guide.docx")

MAX_IMG_WIDTH_IN = 6.3
MAX_IMG_HEIGHT_IN = 8.6

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

HEADING_COLORS = {
    1: RGBColor(0x1E, 0x3A, 0x5F),
    2: RGBColor(0x2E, 0x6D, 0xA4),
    3: RGBColor(0x2E, 0x86, 0xAB),
    4: RGBColor(0x45, 0x8B, 0x74),
}

FIGURE_COUNTER = {'n': 0}


# ════════════════════════════════════════════════════════════════════════
# HELPERS
# ════════════════════════════════════════════════════════════════════════

def set_heading(text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.color.rgb = HEADING_COLORS.get(level, RGBColor(0, 0, 0))
    run.font.bold = True
    run.font.size = Pt({1: 18, 2: 14, 3: 12, 4: 11}.get(level, 11))
    return h


def shade_cell(cell, hex_color='D9EAF7'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], 'D9EAF7')
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            if cells[ci].paragraphs[0].runs:
                cells[ci].paragraphs[0].runs[0].font.size = Pt(9.5)
            if ri % 2 == 0:
                shade_cell(cells[ci], 'F5FAFF')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run('\u2139  ' + text)
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x6E, 0x91)
    p.paragraph_format.left_indent = Inches(0.2)
    return p


def add_warning(text):
    p = doc.add_paragraph()
    run = p.add_run('\u26A0  ' + text)
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xB4, 0x5A, 0x09)
    p.paragraph_format.left_indent = Inches(0.2)
    return p


def _runs(p, text):
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    text = re.sub(r'^[-*]\s*', '', text)
    _runs(p, text)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p


def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    text = re.sub(r'^\d+\.\s*', '', text)
    _runs(p, text)
    return p


def add_para(text):
    p = doc.add_paragraph()
    _runs(p, text)
    return p


def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_image(filename, caption=None, max_w=MAX_IMG_WIDTH_IN, max_h=MAX_IMG_HEIGHT_IN):
    """Insert a screenshot, scaled to fit the page, centered, with an
    optional caption below it. Adds a light border via a 1-cell borderless
    table wrapper is overkill — we just add the picture directly."""
    path = SHOTS / filename
    if not path.exists():
        add_warning(f'[Missing screenshot: {filename}]')
        return
    with Image.open(path) as im:
        px_w, px_h = im.size
    aspect = px_h / px_w
    w_in = max_w
    h_in = w_in * aspect
    if h_in > max_h:
        h_in = max_h
        w_in = h_in / aspect

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(w_in), height=Inches(h_in))

    if caption:
        FIGURE_COUNTER['n'] += 1
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"Figure {FIGURE_COUNTER['n']}. {caption}")
        r.font.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    doc.add_paragraph()


def feature_section(number, title, screenshot, caption, description,
                     fields=None, statuses=None, steps_title=None, steps=None,
                     notes=None, warnings=None, extra_image=None, extra_caption=None):
    set_heading(f'{number} {title}', 2)
    if description:
        add_para(description)
    if screenshot:
        add_image(screenshot, caption)
    if steps:
        if steps_title:
            set_heading(steps_title, 4)
        for s in steps:
            add_numbered(s)
    if extra_image:
        add_image(extra_image, extra_caption)
    if fields:
        set_heading('Key Fields', 4)
        add_table(['Field', 'Description'], fields, col_widths=[1.8, 4.5])
    if statuses:
        set_heading('Statuses', 4)
        add_table(['Status', 'Meaning'], statuses, col_widths=[1.6, 4.7])
    if notes:
        for n in notes:
            add_note(n)
    if warnings:
        for w in warnings:
            add_warning(w)
    doc.add_paragraph()


def add_toc_field():
    """Insert a real Word TOC field (auto-updating on 'Update Field')."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_text = OxmlElement('w:t')
    fld_text.text = "Right-click here and choose 'Update Field' (or press F9) to generate the Table of Contents."
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_sep)
    r_element.append(fld_text)
    r_element.append(fld_end)


print('Building document...')

# ════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_heading('TOS Complete User Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
title.runs[0].font.size = Pt(30)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Tri-E Trading Operating System (TOS)')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)
r.font.bold = True

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2b = sub2.add_run('Tri-E Enterprises — Internal Operations System')
r2b.font.size = Pt(12)
r2b.font.color.rgb = RGBColor(0x44, 0x6E, 0x91)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = desc.add_run('A complete, illustrated walkthrough of every module — Point of Sale, Inventory,\n'
                   'Sales & Purchasing, Delivery & Fleet Management, Finance, HR & Payroll,\n'
                   'Reports, and User Administration.')
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
r2.font.italic = True

doc.add_paragraph()
doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
rm = meta.add_run('For Administrators, Supervisors, Cashiers, and Staff\nDocumentation Version 2.0 · August 2026')
rm.font.size = Pt(10)
rm.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()
doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════
set_heading('Table of Contents', 1)
add_toc_field()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# GETTING STARTED — placeholder marker for later insertion (kept in order)
# ════════════════════════════════════════════════════════════════════════

# ════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════════════════════════════════════
set_heading('1. Introduction', 1)

set_heading('1.1 About Tri-E TOS', 2)
add_para('The **Tri-E Trading Operating System (TOS)** is an all-in-one business management '
         'platform built for Tri-E Enterprises’ trading and retail operations. It brings '
         'together every part of the business into one system:')

for b in [
    '**Point of Sale (POS)** — fast, touch-friendly counter transactions',
    '**Inventory & Sales** — products, categories, customers, suppliers, sales, purchases, quotations',
    '**Delivery Management** — delivery tracking and driver assignment',
    '**Fleet Management** — company vehicles and maintenance scheduling',
    '**Attendance Management** — employee time tracking, leave, and biometric devices',
    '**Payroll** — payroll runs, compensation, and government contributions',
    '**Finance** — expenses, financial dashboard, and profit & loss reporting',
    '**Reports** — sales, inventory, aging, and driver performance analytics',
    '**User Administration** — users, roles & permissions, and audit logs',
]:
    add_bullet(b)

doc.add_paragraph()
set_heading('1.2 System Requirements', 2)
for b in [
    'A modern web browser — Chrome, Edge, Firefox, or Safari',
    'An internet connection (or local network access to the server)',
    'For mobile/tablet POS use: a touch-capable iOS or Android device',
]:
    add_bullet(b)

doc.add_paragraph()
set_heading('1.3 Technology at a Glance', 2)
add_table(
    ['Layer', 'Technology'],
    [
        ['Framework', 'Laravel 12 (PHP 8.2+)'],
        ['Admin Panel', 'Filament v4 with Livewire v3'],
        ['Access Control', 'Spatie Permissions + Filament Shield (role-based)'],
        ['POS Terminal', 'Alpine.js single-page application'],
        ['Database', 'MySQL'],
        ['PDF / Reports', 'Laravel Dompdf, CSV / Excel export'],
        ['Biometric Integration', 'ZKTeco time-clock devices (attendance)'],
    ],
    col_widths=[2.0, 4.3],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 2. GETTING STARTED
# ════════════════════════════════════════════════════════════════════════
set_heading('2. Getting Started', 1)

set_heading('2.1 Logging In', 2)
for s in [
    'Open your browser and go to the TOS URL provided by your administrator.',
    'Enter your **email address** and **password**.',
    'Optionally tick **Remember me** to stay signed in on this device.',
    'Click **Sign in**.',
]:
    add_numbered(s)
add_image('00-login-page.png', 'The Sign in screen.')
add_note('If you forgot your password, contact your system administrator to have it reset.')

set_heading('2.2 The Home Screen (Module Launcher)', 2)
add_para('After logging in, the **Home** page shows a card for every module you have access '
         'to — this is the fastest way to get around the system. Each card opens straight into '
         'that module. Modules you do not have permission for simply will not appear.')
add_image('02-home-launcher.png', 'The Home screen — every module you can access, as a clickable card grid.')
add_note('Access to each card is controlled by your assigned **role** (see Section 12.2, '
         '"Roles & Permissions"). A cashier, for example, will typically only see Point of Sale '
         'and My Attendance.')

set_heading('2.3 Top Navigation Bar', 2)
add_para('In addition to the Home screen, a top navigation bar is always visible with quick links to:')
add_table(
    ['Menu', 'Contains'],
    [
        ['Dashboard', 'The business metrics dashboard (Section 3)'],
        ['POS', 'Opens the Point of Sale terminal directly'],
        ['Reports ▾', 'Inventory In/Out, Products, Sales, Supplier Price Comparison, Daily Transaction, Driver KPI, Aging Report'],
        ['Authentication ▾', 'Roles & Permissions'],
        ['Finance ▾', 'Financial Dashboard, Profit & Loss'],
        ['Search (top right)', 'Global search across the system'],
        ['Avatar (top right)', 'Your Profile and Sign out'],
    ],
    col_widths=[1.6, 4.7],
)

set_heading('2.4 User Roles', 2)
add_para('Access in TOS is role-based. The system ships with these common roles (your '
         'administrator can create additional custom roles at any time):')
add_table(
    ['Role', 'Typical Access'],
    [
        ['Super Admin', 'Full, unrestricted access to every module and setting'],
        ['Admin', 'Broad administrative access across modules'],
        ['Manager / Ops Supervisor', 'Operational modules — sales, inventory, deliveries, reports'],
        ['Cashier', 'Point of Sale, plus basic sales and customer lookup'],
        ['Sales Rep', 'Customers, quotations, and sales'],
        ['Warehouse Clerk', 'Products, categories, purchases, inventory'],
        ['Driver', 'My Attendance and assigned deliveries only'],
    ],
    col_widths=[2.0, 4.3],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 3. DASHBOARD
# ════════════════════════════════════════════════════════════════════════
set_heading('3. Dashboard', 1)
add_para('The **Dashboard** (click **Dashboard** in the top nav) gives an at-a-glance view of '
         'overall business health.')

for b in [
    '**Financial Overview** — Monthly Revenue, Monthly Profit, Monthly Expenses, and Profit '
    'Margin, each compared against last month with a trend indicator.',
    '**Store Overview** — quick totals for Products, Sales, Total Sales Price, Purchases, '
    'and Low Stock items.',
    '**Overdue Alerts** — Accounts Receivable (unpaid customer invoices) and Accounts Payable '
    '(unpaid supplier bills) that have passed their due date, listed individually with days overdue.',
    '**Collection & Payment Reminders** — invoices due today or within the next 7 days, so you '
    'can plan collections proactively.',
]:
    add_bullet(b)

add_image('01-dashboard-widgets.png', 'The Dashboard — financial overview, store stats, and overdue-account alerts.')
add_note('Clicking any invoice or bill line on the dashboard opens that record directly. '
         'Use **View Full Aging Report** to see the complete overdue breakdown (Section 11.7).')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 4. POINT OF SALE (POS)
# ════════════════════════════════════════════════════════════════════════
set_heading('4. Point of Sale (POS)', 1)
add_para('The POS is a fast, touch-friendly, single-page terminal for processing counter sales. '
         'It works on desktop, tablet, and mobile, and runs without full-page reloads — a '
         'session can stay open all day. Open it from the **POS** link in the top nav, the '
         '**Point of Sale** card on the Home screen, or by browsing directly to `/pos`.')
add_image('03-pos.png', 'The POS terminal — product grid on the left, current sale on the right.')

set_heading('4.1 Opening the Cash Register', 2)
add_para('You must **open the register** before any sale can be processed.')
for s in [
    'When you first open the POS, you will be prompted for an **Opening Amount**.',
    'Count the physical cash in the drawer and enter the total.',
    'Click **Open Register**.',
]:
    add_numbered(s)
add_note('Only one open register session per user is allowed at a time. The register bar at the '
         'top of the POS shows a green **Register Open** badge once active.')

set_heading('4.2 Browsing and Adding Products', 2)
for b in [
    'Browse **product cards**, sorted by stock level, or use the **search bar** to find a product by name.',
    'Use the **category dropdown** to narrow the list.',
    'Click a product card to add it — for weight, volume, or length units (kg, g, L, mL, m, ft) a '
    'quantity modal appears first; all other units are added directly with a quantity of 1.',
    'Products with zero stock are marked out-of-stock and cannot be added.',
    'Use the green **+ Custom** button to add a manual line item (name, unit, price, quantity) '
    'for anything not in the catalogue — manual items do not affect inventory.',
]:
    add_bullet(b)

set_heading('4.3 Managing the Cart', 2)
add_table(
    ['Action', 'How'],
    [
        ['Increase / decrease quantity', 'Use the + / − buttons on the cart line'],
        ['Type a quantity directly', 'Click the quantity number and type a new value'],
        ['Apply a discount', 'Enter a ₱ amount on the line; tick "flat total" to make it a fixed '
                              'discount for the whole line instead of per piece'],
        ['Remove an item', 'Click the trash / × icon on that line'],
        ['Clear the cart', 'Click Clear Cart — a confirmation prompt appears first'],
    ],
    col_widths=[2.1, 4.2],
)
add_note('The cart total updates live. A discount can never make a line total go below zero.')

set_heading('4.4 Selecting a Customer', 2)
for b in [
    'Search for an existing customer by name in the **Customer** field, or click the **+** icon '
    'to add a new customer on the spot (name, phone, email, address) — the new customer is '
    'auto-selected after saving.',
    'Leave the field on **Walk-in Customer** for anonymous sales.',
]:
    add_bullet(b)

set_heading('4.5 Processing a Payment', 2)
for s in [
    'Review the cart total, then click **Charge** (Complete Sale).',
    'Select a **payment method**: Cash, Bank Transfer, Check, Credit Card, GCash, or Maya.',
    'For **Cash**, enter the amount received — the system calculates the change automatically.',
    'Optionally add a flat **Delivery Fee**, applied on top of the discounted item total.',
    'Optionally toggle **Add Payment Terms** for a credit sale (see 4.6 below).',
    'Click **Confirm Payment**.',
]:
    add_numbered(s)
add_para('The payment modal breaks the total down as:')
add_table(
    ['Line', 'Meaning'],
    [
        ['Items Total', 'Raw subtotal — unit price × quantity, before discounts'],
        ['− Discount', 'Sum of all per-item discounts (only shown if greater than zero)'],
        ['+ Delivery Fee', 'Optional flat fee (only shown if greater than zero)'],
        ['= Grand Total', 'What the customer actually pays'],
    ],
    col_widths=[1.8, 4.5],
)
add_note('On confirmation: inventory is decremented for every non-manual item, the register '
         'totals update (for immediately-paid sales only), and a success screen with receipt '
         'print options appears.')

set_heading('4.6 Payment Terms (Credit Sales)', 2)
add_para('Use payment terms **only** for trusted customers with an approved credit arrangement.')
for s in [
    'In the payment modal, toggle on **Add Payment Terms**.',
    'Select the credit period: **5, 10, 15, 30, or 60 days**.',
    'The **Due Date** is calculated automatically.',
    'Choose a payment method and click **Confirm Payment**.',
]:
    add_numbered(s)
add_table(
    ['Condition', 'Payment Status', 'Counted in Register Totals?'],
    [
        ['No payment terms (cash, card, GCash, etc.)', 'paid', 'Yes'],
        ['Payment terms applied', 'unpaid', 'No — until collected'],
    ],
    col_widths=[2.6, 1.6, 2.1],
)
add_warning('Never apply payment terms to an instant/cash payment. Terms are only for genuine '
            'credit arrangements — the sale is excluded from cash totals until it is marked paid '
            '(Section 5.5).')

set_heading('4.7 Printing Receipts', 2)
add_table(
    ['Receipt Type', 'Use When'],
    [
        ['Delivery Receipt', 'The order will be delivered — includes driver / receiver signature lines'],
        ['Pick Up Receipt', 'The customer collects at the counter — includes released-by / picked-up-by signature lines'],
    ],
    col_widths=[1.8, 4.5],
)
add_note('Receipts print two copies side by side on A4 landscape — an Office Copy and a '
         "Customer's Copy — separated by a cut line.")

set_heading('4.8 Reprinting a Past Receipt', 2)
add_para('Click **Reprint** in the POS header to see the last 50 sales, most recent first. '
         'Search by receipt number or customer name, then reprint in either format.')

set_heading('4.9 Creating a Quotation', 2)
for s in [
    'Build the cart as normal and optionally select a customer.',
    'Click **Create Quotation** instead of Complete Sale.',
    'Set the **validity period** (7–90 days, default 30) and add optional notes.',
    'Click **Create & Print** — a printable quotation opens in a new window.',
]:
    add_numbered(s)
add_para('Once created, an admin with the `approve_quotation` permission is notified by email '
         'and can approve or reject it. An **approved** quotation can be converted back into a '
         'sale by re-opening the POS from the quotation link — the cart is pre-filled with its items.')
add_table(
    ['Status', 'Meaning'],
    [
        ['Pending', 'Awaiting admin approval'],
        ['Approved', 'Ready to convert to a sale'],
        ['Rejected', 'Declined by admin'],
        ['Converted to Sale', 'Already completed as a sale'],
        ['Expired', 'Validity period has passed'],
    ],
    col_widths=[1.8, 4.5],
)

set_heading('4.10 Settling an Invoice', 2)
add_para('Use **Settle Invoice** in the POS header to record a payment against an existing '
         'outstanding (credit/term) sale without creating a new transaction, and print a '
         'payment receipt.')

set_heading('4.11 Closing the Cash Register', 2)
for s in [
    'Click **Close Register**.',
    'Physically count all cash in the drawer and enter the **Closing Amount**.',
    'Add any notes (e.g., unusual transactions) if needed.',
    'Click **Close Register** to confirm.',
]:
    add_numbered(s)
add_para('The system calculates the **expected amount** (opening cash + cash sales collected) '
         'and shows the **discrepancy** (positive = overage, negative = shortage). A PDF '
         'session report downloads automatically, containing the summary and a full list of '
         'every sale processed in the session.')
add_note('Re-download a past closure report anytime from the Cash Register list (Section 5.8).')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 5. INVENTORY & SALES
# ════════════════════════════════════════════════════════════════════════
set_heading('5. Inventory & Sales', 1)
add_para('This module group covers the product catalogue, customers, suppliers, and every '
         'sales and purchasing transaction. Open any of these from the **Home** screen.')

feature_section(
    '5.1', 'Products', '19-products.png', 'The Products list — searchable, with category, price, stock, and unit.',
    'The product catalogue is the source of truth for everything sold through the POS.',
    steps_title='Adding a Product',
    steps=[
        'Click **New product**.',
        'Fill in **Name**, **Category**, **Supplier**, **Price** (selling price), and **Stock**.',
        'Choose the **Unit** (see the unit table below).',
        'Optionally configure **Additional Sellable Units** (e.g. sell a hose per Meter but also '
        'per Roll) and **Supplier Base Prices** (record what each supplier charges, for price '
        'comparison on purchase orders).',
        'Click **Create**.',
    ],
    extra_image='50-product-create-form.png',
    extra_caption='The Create Product form.',
    fields=[
        ['Price / Cost Price', 'Selling price and purchase cost — used to calculate profit margin automatically'],
        ['Unit', 'Piece, Bag, Box, Bundle, Knot, Tube, Kilo, Gram, Liter, Milliliter, Meter, Foot, Cubic Meter'],
        ['Profit Margin / Profit per Unit', 'Calculated automatically from Price and Cost Price (defaults to 70% margin if cost is not set)'],
    ],
    notes=['Stock levels are updated automatically through **Purchases** (Section 5.6) — do not edit stock manually.'],
)

feature_section(
    '5.2', 'Categories', '20-categories.png', 'The Categories list.',
    'Categories organise products for easier browsing in the POS product grid and reports. '
    'Simply add a **Name** for each category you need.',
)

feature_section(
    '5.3', 'Suppliers', '21-suppliers.png', 'The Suppliers list.',
    'Suppliers are the vendors you purchase stock from.',
    fields=[
        ['Contact Person / Phone / Email / Address', 'Supplier contact details'],
        ['Payment Term Days', '0 = COD (cash on delivery); otherwise Net X days — used to auto-calculate the due date on purchases from this supplier'],
    ],
)

feature_section(
    '5.4', 'Customers', '18-customers.png', 'The Customers list.',
    'Customer profiles are shared across Sales, Quotations, and Deliveries.',
    steps_title='Adding a Customer',
    steps=[
        'Click **New customer**.',
        'Enter **Name** (required) plus Company, Phone, Email, and Address as available.',
        'Set **Payment Term Days** if this customer is allowed to pay on credit.',
        'Click **Create**.',
    ],
    extra_image='51-customer-create-form.png',
    extra_caption='The Create Customer form.',
    notes=['Customers can also be added on the fly directly from the POS without leaving the terminal (Section 4.4).'],
)

feature_section(
    '5.5', 'Sales', '17-sales.png', 'The Sales list — customer, date, item count, payment status, and total.',
    'Every completed transaction — from the POS or created manually — appears here.',
    statuses=[
        ['paid (green)', 'Payment has been fully collected'],
        ['partial (yellow)', 'Part of the amount has been collected'],
        ['unpaid (red)', 'No payment collected yet — a credit/term sale'],
    ],
    steps_title='Creating a Manual Sale',
    steps=[
        'Click **New sale**.',
        'Select a customer (or leave blank for walk-in).',
        'Add products and quantities, and set the payment status.',
        'Click **Create**.',
    ],
    extra_image='52-sale-create-form.png',
    extra_caption='The Create Sale form.',
    notes=[
        'Most sales should be processed through the POS for the full discount, payment, and receipt experience.',
        'A **Mark as Paid** action appears on unpaid credit/term sales — click it, confirm the amount, and the sale is marked paid immediately.',
        'Use **Download Summary** to export a CSV of sales totals grouped by date, for any period (Today, This Week, This Month, custom range, etc.).',
    ],
)

feature_section(
    '5.6', 'Purchases', '22-purchases.png', 'The Purchases list.',
    'Purchases record stock received from suppliers and automatically update inventory.',
    steps_title='Creating a Purchase',
    steps=[
        'Click **New purchase**.',
        'Select the **Supplier**.',
        'Add purchase items — product, quantity ordered, and unit price.',
        'Click **Create**.',
    ],
    extra_image='53-purchase-create-form.png',
    extra_caption='The Create Purchase form.',
    statuses=[
        ['Pending', 'No items received yet'],
        ['Partial', 'At least one item received, but not all'],
        ['Received', 'All ordered items fully received'],
    ],
    notes=[
        'When a supplier can only fulfil part of an order, record the actual **quantity received** '
        'per line — inventory increases by exactly that amount. The amount owed is always based '
        'on the full ordered quantity, regardless of how much has arrived.',
        'The due date is calculated automatically from the supplier’s payment terms.',
    ],
)

feature_section(
    '5.7', 'Quotations', '16-quotations.png', 'The Quotations list.',
    'Quotations are price proposals sent to a customer before a sale is confirmed — created '
    'either here or from the POS (Section 4.9). Open a quotation to **Approve**, **Reject**, or '
    'print it; an approved quotation can be converted into a sale from the POS.',
)

feature_section(
    '5.8', 'Cash Register', '39-cash-register.png', 'The Cash Register Sessions list.',
    'A full history of every cash register session — opening amount, closing amount, total '
    'sales, and any discrepancy. Open a session to re-download its PDF closure report.',
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 6. DELIVERY MANAGEMENT
# ════════════════════════════════════════════════════════════════════════
set_heading('6. Delivery Management', 1)
add_para('Track the delivery of orders to customers, from assignment to drop-off.')

feature_section(
    '6.1', 'Deliveries', '23-deliveries.png', 'The Deliveries list.',
    'A delivery is linked to a completed sale.',
    steps_title='Creating a Delivery',
    steps=[
        'Click **New delivery**.',
        'Select the **Sale** this delivery is for and assign a **Driver**.',
        'Enter the delivery **Address**, distance, and any notes.',
        'Click **Create**.',
    ],
    extra_image='55-delivery-create-form.png',
    extra_caption='The Create Delivery form.',
    statuses=[
        ['Pending', 'Delivery created, awaiting driver assignment'],
        ['Assigned', 'Driver has been assigned'],
        ['Picked Up', 'Driver has collected the goods'],
        ['In Transit', 'Goods are on the way'],
        ['Delivered', 'Successfully delivered to the customer'],
        ['Failed', 'Delivery attempt was unsuccessful'],
        ['Returned', 'Goods were returned to the warehouse'],
    ],
    notes=['Open a delivery and use **Print Receipt** to generate the driver’s delivery '
           'document with signature lines. Customers can also leave a 1–5 star rating and feedback.'],
)

feature_section(
    '6.2', 'Drivers', '24-drivers.png', 'The Drivers list.',
    'Manage delivery drivers and see their performance at a glance.',
    fields=[
        ['Name / Phone / License Number', 'Driver identification and contact details'],
        ['Active Status', 'Inactive drivers are hidden from delivery assignment'],
        ['Delivery Count / Average Rating / Success Rate', 'Calculated automatically from completed deliveries'],
    ],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 7. FLEET MANAGEMENT
# ════════════════════════════════════════════════════════════════════════
set_heading('7. Fleet Management', 1)
add_para('Track and maintain the company vehicle fleet, from acquisition to scheduled servicing.')

feature_section(
    '7.1', 'Vehicles', '25-vehicles.png', 'The Vehicles list.',
    'Every company vehicle, with its maintenance history rolled up automatically.',
    steps_title='Adding a Vehicle',
    steps=[
        'Click **New vehicle**.',
        'Enter the **Plate Number** (unique), Make, Model, Year, VIN, and Engine Number.',
        'Set Fuel Type, Transmission, Current Mileage, and optionally assign a Driver.',
        'Click **Create**.',
    ],
    extra_image='56-vehicle-create-form.png',
    extra_caption='The Create Vehicle form.',
    fields=[
        ['Status', 'Active / Maintenance / Inactive / Sold'],
        ['Total Maintenance Cost / Last Maintenance Date', 'Calculated automatically from service records'],
        ['Maintenance Due', 'Flagged automatically when mileage or time exceeds the next-service threshold'],
    ],
)

feature_section(
    '7.2', 'Maintenance Types (Service Types)', '28-maintenance-types.png', 'The Maintenance Types list.',
    'Define the categories of maintenance work performed on vehicles — e.g. Oil Change, Tire '
    'Rotation, Brake Service — including the recommended interval in kilometres and/or months.',
)

feature_section(
    '7.3', 'Maintenance Records (Service Records)', '26-maintenance-records.png', 'The Maintenance Records list.',
    'The complete history of completed vehicle maintenance.',
    fields=[
        ['Parts Cost + Labor Cost', 'Combine into the record’s Total Cost'],
        ['Next Service Date / Next Service Mileage', 'Used to flag the vehicle as due for maintenance'],
        ['Invoice Upload', 'Attach the service provider’s invoice'],
    ],
)

feature_section(
    '7.4', 'Maintenance Requests (Service Requests)', '27-maintenance-requests.png', 'The Maintenance Requests list.',
    'Staff submit a service request for a vehicle; it goes through an approval workflow before '
    'becoming a Maintenance Record.',
    steps_title='Submitting a Request',
    steps=[
        'Click **New request**.',
        'Select the **Vehicle** and **Maintenance Type**.',
        'Describe the issue and set a **Priority** (Low / Normal / High).',
        'Click **Create**.',
    ],
    statuses=[
        ['Pending', 'Awaiting review'],
        ['Approved', 'Approved, with an estimated cost — ready to be actioned'],
        ['Rejected', 'Declined, with a rejection reason recorded'],
        ['Completed', 'Linked to the resulting Maintenance Record'],
    ],
    notes=['A badge on the Home card shows the current count of pending requests awaiting approval.'],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 8. ATTENDANCE MANAGEMENT
# ════════════════════════════════════════════════════════════════════════
set_heading('8. Attendance Management', 1)
add_para('Employee time tracking, leave management, and biometric device integration.')

feature_section(
    '8.1', 'My Attendance', '15-my-attendance.png', 'The My Attendance self-service page.',
    'Every employee has access to their own **My Attendance** page to clock in, clock out, and '
    'review their attendance history — no admin access required.',
)

feature_section(
    '8.2', 'Attendance', '31-attendance.png', 'The company-wide Attendance list.',
    'Administrators and supervisors can view and manage attendance logs for every employee here.',
    statuses=[
        ['Present (green)', 'Clocked in and out as expected'],
        ['Absent (red)', 'No attendance recorded for the day'],
        ['Late (yellow)', 'Clocked in after the expected time'],
        ['Half Day (blue)', 'Worked only part of the day'],
        ['On Leave (gray)', 'Covered by an approved leave request'],
    ],
    notes=['**Total Hours** is calculated automatically from Time In and Time Out. Filter by '
           'status, employee, or date range.'],
)

feature_section(
    '8.3', 'Employees', '30-employees.png', 'The Employees list.',
    'Employee records — every system user can also be an employee, linked to their attendance, '
    'leave, compensation, and payroll history.',
    steps_title='Adding an Employee',
    steps=[
        'Click **New employee**.',
        'Enter the employee’s details, including their **biometric PIN** (for ZKTeco devices) '
        'and login credentials if they need system access.',
        'Click **Create**.',
    ],
    extra_image='57-employee-create-form.png',
    extra_caption='The Create Employee form.',
)

feature_section(
    '8.4', 'Biometric Devices', '32-biometric-devices.png', 'The Biometric Devices list.',
    'Register and monitor ZKTeco biometric time-clock terminals. Devices push attendance '
    'events (finger-print or face clock-ins) directly into the Attendance module — no manual '
    'entry required at sites with a registered device.',
)

feature_section(
    '8.5', 'Leave Types', '34-leave-types.png', 'The Leave Types list.',
    'Configure the leave categories available to employees (e.g. Sick Leave, Vacation Leave), '
    'each with an annual day entitlement and whether it is paid.',
)

feature_section(
    '8.6', 'Leave Requests', '33-leave-requests.png', 'The Leave Requests list.',
    'Employees submit leave requests, which go through an approval workflow.',
    statuses=[
        ['Pending', 'Awaiting a decision'],
        ['Approved', 'Approved — total days are deducted from the entitlement'],
        ['Rejected', 'Declined, with a reason recorded'],
        ['Cancelled', 'Withdrawn by the employee or an admin'],
    ],
    notes=['A badge on the Home card shows the current count of pending leave requests.'],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 9. PAYROLL
# ════════════════════════════════════════════════════════════════════════
set_heading('9. Payroll', 1)
add_para('Process payroll runs and manage the compensation setup behind them.')

feature_section(
    '9.1', 'Payrolls', '29-payroll.png', 'The Payrolls list.',
    'Each payroll run covers a pay period (Daily, Weekly, or Semi-Monthly) and follows a '
    '**Draft → Approved → Paid** workflow. Individual employee pay lines (gross pay, '
    'deductions, net pay) are managed inside each payroll record.',
    statuses=[
        ['Draft', 'Being prepared — can still be edited'],
        ['Approved', 'Reviewed and approved — ready to be paid'],
        ['Paid', 'Payment has been released'],
        ['Cancelled', 'Voided — not possible once Paid'],
    ],
    notes=['A badge on the Home card shows the count of draft payrolls awaiting approval. Each '
           'payroll can be exported as a PDF payslip.'],
)

feature_section(
    '9.2', 'Employee Compensation', '35-compensations.png', 'The Employee Compensation list.',
    'The salary and compensation setup for each employee — the source figures a payroll run '
    'pulls from when calculating gross pay.',
)

feature_section(
    '9.3', 'Government Contributions', '36-gov-contributions.png', 'The Government Contributions list.',
    'Configure statutory deduction tables — SSS, PhilHealth, and Pag-IBIG — used to calculate '
    'mandatory deductions during payroll processing.',
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 10. FINANCE
# ════════════════════════════════════════════════════════════════════════
set_heading('10. Finance', 1)
add_para('Track spending, and monitor overall financial performance.')

feature_section(
    '10.1', 'Expenses', '37-expenses.png', 'The Expenses list.',
    'Record every business expense — utilities, rent, supplies, vehicle maintenance, salaries, and more.',
    steps_title='Recording an Expense',
    steps=[
        'Click **New expense**.',
        'Select the **Category**, enter the **Amount**, **Payee**, and a **Description**.',
        'Set the **Expense Date** (cannot be a future date) and payment method.',
        'Optionally attach a **Receipt** (image or PDF, max 5 MB).',
        'Click **Create**.',
    ],
    extra_image='54-expense-create-form.png',
    extra_caption='The Create Expense form.',
    fields=[
        ['Reference Number', 'Auto-generated, format EXP-YYYYMMDD-XXXX'],
        ['Status', 'pending / approved / rejected (defaults to approved)'],
    ],
    notes=['Only **approved** expenses are counted in financial totals and the Profit & Loss report.'],
)

feature_section(
    '10.2', 'Expense Categories', '38-expense-categories.png', 'The Expense Categories list.',
    'Create and manage expense categories (Utilities, Rent, Supplies, Transportation, Salaries, '
    'Marketing, Maintenance, etc.) to keep spending organised. Each category shows its running total.',
)

feature_section(
    '10.3', 'Financial Dashboard', '12-financial-dashboard.png', 'The Financial Dashboard.',
    'A comprehensive financial overview for any period.',
    fields=[
        ['Total Revenue', 'Income from sales in the selected period'],
        ['Total Expenses', 'All approved expenses in the period'],
        ['Gross Profit', 'Revenue minus cost of goods sold (COGS)'],
        ['Net Profit', 'Gross profit minus operating expenses'],
    ],
    steps_title='Choosing a Period',
    steps=[
        'Select **Today, This Week, This Month, This Year**, or a **Custom Date Range**.',
        'Review the expense breakdown by category.',
        'Use **Export to PDF** or **Export to Excel** to download the report.',
    ],
)

feature_section(
    '10.4', 'Profit & Loss', '13-profit-loss.png', 'The Profit & Loss report.',
    'A detailed profit and loss statement: revenue breakdown, cost of goods sold, gross margin, '
    'operating expenses by category, and net income for the selected period.',
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 11. REPORTS
# ════════════════════════════════════════════════════════════════════════
set_heading('11. Reports', 1)
add_para('All analytical and export-focused reports live under **Reports** in the top nav or '
          'as cards on the Home screen. Most reports support a date-range filter and CSV / '
          'Excel / PDF export.')

feature_section(
    '11.1', 'Daily Transaction Report', '08-daily-transaction-report.png', 'The Daily Transaction Report.',
    'A day-by-day breakdown of transactions processed through the POS — useful for a quick '
    'daily reconciliation.',
)

feature_section(
    '11.2', 'Sales Report', '06-sales-report.png', 'The Sales Report.',
    'Comprehensive sales analysis filterable by period (Today, Week, Month, Year, or a custom '
    'range) and viewable by customer. Export to CSV for further analysis.',
)

feature_section(
    '11.3', 'Inventory In/Out Report', '04-inventory-in-out.png', 'The Inventory In/Out Report.',
    'Tracks every inventory movement — stock added from purchases, stock sold, and manual '
    'adjustments — exportable to CSV.',
)

feature_section(
    '11.4', 'Products Report', '05-products-report.png', 'The Products Report.',
    'Product performance analysis, filterable by category and supplier, with current stock levels.',
)

feature_section(
    '11.5', 'Supplier Price Comparison', '07-supplier-price-comparison.png', 'The Supplier Price Comparison Report.',
    'Compares what each supplier charges for the same product, using the base prices recorded '
    'on the product record (Section 5.1) — helpful when deciding who to purchase from.',
)

feature_section(
    '11.6', 'Driver KPI Dashboard', '09-driver-kpi.png', 'The Driver KPI Dashboard.',
    'Delivery driver performance at a glance.',
    fields=[
        ['Total / Successful / Failed Deliveries', 'Delivery counts by outcome'],
        ['On-time Rate', 'Percentage of deliveries completed on schedule'],
        ['Average Delivery Time', 'Average time from pickup to delivery'],
    ],
)

feature_section(
    '11.7', 'Aging Report', '10-aging-report.png', 'The Aging Report.',
    'Tracks overdue payments in both directions.',
    fields=[
        ['Accounts Receivable', 'Money owed TO you by customers'],
        ['Accounts Payable', 'Money YOU owe to suppliers'],
    ],
    statuses=[
        ['Current', 'Not yet due'],
        ['1–30 Days', 'Overdue by 1 to 30 days'],
        ['31–60 Days', 'Overdue by 31 to 60 days'],
        ['61–90 Days', 'Overdue by 61 to 90 days'],
        ['Over 90 Days', 'Overdue by more than 90 days'],
    ],
    notes=['Export to Excel or PDF for follow-up with customers or suppliers.'],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 12. AUTHENTICATION & SECURITY
# ════════════════════════════════════════════════════════════════════════
set_heading('12. Authentication & Security', 1)
add_para('User accounts, access control, and a full audit trail of system activity.')

feature_section(
    '12.1', 'Users', '40-user-management.png', 'The User Management list.',
    'Manage system users — add new users, edit user information, reset passwords, and assign roles.',
    notes=['Every user also functions as an employee, linked to Attendance, Leave Requests, '
           'Compensation, and Payroll records (Section 8 and 9).'],
)

feature_section(
    '12.2', 'Roles & Permissions', '11-roles.png', 'The Roles & Permissions list.',
    'Roles control exactly which modules and actions each user can access. Every module '
    'exposes granular View / Create / Edit / Delete permissions that can be assigned to any role.',
    notes=['A notable custom permission, **approve_quotation**, controls who can approve or '
           'reject quotations and receive the related email notifications (Section 4.9).'],
)

feature_section(
    '12.3', 'Audit Logs', '41-audit-logs.png', 'The Audit Logs list.',
    'A complete, read-only trail of significant actions across the system — who did what, and '
    'when.',
    fields=[
        ['Action', 'e.g. completed_sale, approved, opened_register'],
        ['Old Values / New Values', 'The record’s field values before and after the change'],
        ['IP Address / User Agent', 'Where the action came from'],
    ],
    notes=['Audit logs cannot be edited or deleted from the interface, and sensitive fields '
           '(passwords, tokens) are always excluded.'],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 13. TIPS & BEST PRACTICES
# ════════════════════════════════════════════════════════════════════════
set_heading('13. Tips & Best Practices', 1)

set_heading('13.1 Daily Operations', 2)
add_para('**Start of day:** log in, check the Dashboard for alerts, review overdue payments '
         '(Aging Report), and check low-stock items.')
add_para('**During the day:** process sales through the POS, record expenses immediately, and '
         'keep delivery statuses up to date.')
add_para('**End of day:** review daily sales (Sales Report / Daily Transaction Report), confirm '
         'all deliveries are updated, and close the cash register.')

set_heading('13.2 Inventory', 2)
for b in [
    'Record every stock addition as a **Purchase** — never edit stock levels manually.',
    'Check physical stock against the system regularly.',
    'Keep products categorised for faster POS browsing.',
]:
    add_bullet(b)

set_heading('13.3 Finance', 2)
for b in [
    'Record expenses daily — do not let them pile up.',
    'Use the correct expense category for accurate Profit & Loss reporting.',
    'Review the Financial Dashboard and P&L weekly.',
    'Follow up on receivables using the Aging Report.',
]:
    add_bullet(b)

set_heading('13.4 Customers & Credit', 2)
for b in [
    'Collect full contact details for every customer.',
    'Set payment terms explicitly and be clear about payment expectations.',
    'Use Quotations for large orders before committing to a sale.',
    'Never apply payment terms to a cash sale.',
]:
    add_bullet(b)

set_heading('13.5 Fleet, Delivery & HR', 2)
for b in [
    'Follow recommended maintenance intervals — the system flags overdue vehicles automatically.',
    'Keep vehicle mileage updated after every trip or service.',
    'Use the Driver KPI dashboard for accountability and route planning.',
    'Approve leave requests and maintenance requests promptly to avoid a growing backlog.',
]:
    add_bullet(b)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# 14. APPENDIX — REFERENCE NUMBER FORMATS
# ════════════════════════════════════════════════════════════════════════
set_heading('14. Appendix — Reference Number Formats', 1)
add_para('Every auto-generated document number follows the same pattern and resets daily.')
add_table(
    ['Record', 'Format', 'Example'],
    [
        ['Quotation', 'QT-YYYYMMDD-XXXX', 'QT-20260220-0001'],
        ['Expense', 'EXP-YYYYMMDD-XXXX', 'EXP-20260220-0001'],
        ['Maintenance Record', 'MNT-YYYYMMDD-XXXX', 'MNT-20260220-0001'],
        ['Maintenance Request', 'REQ-YYYYMMDD-XXXX', 'REQ-20260220-0001'],
        ['Leave Request', 'LR-YYYYMMDD-XXXX', 'LR-20260220-0001'],
        ['Payroll', 'PAY-YYYYMMDD-XXXX', 'PAY-20260220-0001'],
    ],
    col_widths=[2.2, 2.4, 1.7],
)

set_heading('Getting Help', 2)
for s in [
    'Check this user guide first.',
    'Contact your system administrator.',
    'Report bugs or issues through the proper support channel.',
]:
    add_numbered(s)

doc.add_paragraph()
divider()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot.add_run('Tri-E Trading Operating System (TOS) — Complete User Guide\n'
                   'Built with Laravel 12, Filament v4, and Livewire v3\n'
                   'Documentation Version 2.0 · August 2026')
fr.font.size = Pt(9)
fr.font.italic = True
fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT))
print(f'Saved: {OUTPUT}')
