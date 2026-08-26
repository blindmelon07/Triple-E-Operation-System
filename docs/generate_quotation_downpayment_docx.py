"""
Generates: Tri-E Enterprises - Quotation Down Payment User Guide.docx

A step-by-step user guide for staff on recording an optional down payment
on a quotation — for a walk-in customer at the POS, and for an online
transaction entered through the back-office admin panel.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

ASSETS = os.path.join(os.path.dirname(__file__), '_downpayment_assets')

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Default paragraph style ───────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ── Heading colours ───────────────────────────────────────────────────────────
HEADING_COLORS = {
    1: RGBColor(0x1E, 0x3A, 0x5F),   # dark navy
    2: RGBColor(0x2E, 0x6D, 0xA4),   # medium blue
    3: RGBColor(0x2E, 0x86, 0xAB),   # steel blue
    4: RGBColor(0x45, 0x8B, 0x74),   # teal green
}

def set_heading(text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.color.rgb = HEADING_COLORS.get(level, RGBColor(0, 0, 0))
    run.font.bold = True
    run.font.size = Pt({1: 18, 2: 14, 3: 12, 4: 11}.get(level, 11))
    return h

def shade_cell(cell, hex_color='D9EAF7'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], 'D9EAF7')
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                shade_cell(cells[ci], 'F5FAFF')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return t

def add_code_block(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F0F0F0')
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_note(text, color=RGBColor(0x44, 0x6E, 0x91), icon='ℹ'):
    p = doc.add_paragraph()
    run = p.add_run(icon + '  ' + text)
    run.font.italic = True
    run.font.size   = Pt(10)
    run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.2)
    return p

def add_warning(text):
    return add_note(text, color=RGBColor(0xA6, 0x3D, 0x2E), icon='⚠')

def add_success(text):
    return add_note(text, color=RGBColor(0x2E, 0x7D, 0x4F), icon='✔')

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    text = re.sub(r'^[-*]\s*', '', text)
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    return p

def add_para(text):
    p = doc.add_paragraph()
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    return p

def add_image(filename, width_in=5.5, caption=None):
    path = os.path.join(ASSETS, filename)
    if not os.path.exists(path):
        add_warning(f'Screenshot not found: {filename}')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cr = cap.add_run(caption)
        cr.font.name = 'Courier New'
        cr.font.size = Pt(8.5)
        cr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        cap.paragraph_format.space_after = Pt(10)

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Quotation Down Payment', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
title.runs[0].font.size = Pt(26)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('User Guide')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = meta.add_run('Tri-E Enterprises · TOS System · Quotations Module')
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = meta2.add_run('Prepared: August 26, 2026')
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
set_heading('Overview', 1)
add_para('When you create a quotation, you can now record a **down payment** — an amount the customer has already paid toward it. This is completely optional. Leave it blank and the quotation works exactly as before.')
doc.add_paragraph()
add_para('When you do enter an amount, the system automatically works out the **Balance Due** (what the customer still owes) and shows it on:')
add_bullet('The quotation screen, while you\'re creating it')
add_bullet('The printed quotation you hand to the customer')
add_bullet('The quotation record in the admin panel')
doc.add_paragraph()
add_note('You can add a down payment whether the customer is a **walk-in** at the counter or the quotation is being entered as an **online transaction** through the admin panel. Both work the same way.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 1. WALK-IN CUSTOMERS (POS)
# ════════════════════════════════════════════════════════════════════════════
set_heading('1. Walk-in Customers (POS)', 1)
add_para('Use this when a customer is at the counter and you\'re building their quotation from the cart.')
doc.add_paragraph()

add_numbered('Add the customer\'s items to the cart as usual.')
add_numbered('Click **Create Quotation**.')
add_numbered('In the window that opens, set **Valid For (Days)** if you want a different validity period than the default.')
add_numbered('If the customer is paying something now, enter the amount in **Down Payment (Optional)**.')
add_numbered('The **Balance Due** below it updates automatically — that\'s what the customer will still owe.')
add_numbered('Add any **Notes**, then click **Create & Print**.')
doc.add_paragraph()

add_image('pos_modal_crop.png', width_in=3.6,
          caption='Example: ₱79.00 total, ₱30.00 down payment entered → ₱49.00 balance due')

add_note('You can\'t enter more than the total — if you type a bigger number, it\'s automatically capped at the quotation total.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 2. ONLINE TRANSACTIONS (ADMIN PANEL)
# ════════════════════════════════════════════════════════════════════════════
set_heading('2. Online Transactions (Admin Panel)', 1)
add_para('Use this when you\'re entering a quotation directly in the back office rather than through the POS cart — for example, for an online inquiry or a phone order.')
doc.add_paragraph()

add_numbered('Go to **Inventory & Sales → Quotations** and click **Create** (or open an existing quotation and click **Edit**).')
add_numbered('Fill in the customer, date, and items as usual.')
add_numbered('In the **Quotation Summary** section, enter an amount in **Down Payment (Optional)** if the customer has already paid something.')
add_numbered('The **Balance Due** next to **Total Amount** updates automatically.')
add_numbered('Click **Save changes**.')
doc.add_paragraph()

add_image('admin_summary.png', width_in=4.3,
          caption='Quotations → Edit → Quotation Summary')

add_note('Just like in the POS, the down payment can\'t exceed the quotation total — it\'s automatically capped when you save.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 3. WHAT THE CUSTOMER SEES ON THE PRINTED QUOTATION
# ════════════════════════════════════════════════════════════════════════════
set_heading('3. What the Customer Sees on the Printed Quotation', 1)
add_para('If a down payment was entered, the printed quotation automatically shows two extra lines under the total — **Down Payment** and **Balance Due** — so the customer has a clear record of what they\'ve paid and what\'s left.')
doc.add_paragraph()
add_para('If no down payment was entered, the printed quotation looks exactly as it always has — no extra lines are added.')
doc.add_paragraph()

add_image('print_view.png', width_in=5.5,
          caption='Printed quotation with a down payment recorded')

# ════════════════════════════════════════════════════════════════════════════
# 4. GOOD TO KNOW
# ════════════════════════════════════════════════════════════════════════════
set_heading('4. Good to Know', 1)
add_bullet('**It\'s always optional.** Skip it and the quotation behaves exactly as before — no down payment, no balance shown.')
add_bullet('**It can\'t exceed the total.** If you enter more than the quotation total, it\'s automatically reduced to match.')
add_bullet('**It carries over on conversion.** If an approved quotation with a down payment is converted into a sale at the POS, the down payment amount is already filled in for you — no need to re-enter it.')
add_bullet('**Existing quotations aren\'t affected.** Any quotation created before this feature shows a down payment of ₱0.00 and a balance equal to the full total.')
doc.add_paragraph()

foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot.add_run('Tri-E Enterprises · TOS System')
fr.font.size = Pt(8.5)
fr.italic = True
fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

OUT_PATH = os.path.join(os.path.dirname(__file__), 'Tri-E Enterprises - Quotation Down Payment User Guide.docx')
doc.save(OUT_PATH)
print('Saved:', OUT_PATH)
