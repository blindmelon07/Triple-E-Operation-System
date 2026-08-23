"""
Generates: Tri-E Enterprises - Inventory Report Fix Documentation.docx

Documents the investigation and fix of the Inventory In/Out Report issue
(missing "in" entries for purchases and manual stock edits).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Default paragraph style ───────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ── Heading colours ───────────────────────────────────────────────────────────
HEADING_COLORS = {
    1: RGBColor(0x1E, 0x3A, 0x5F),   # dark navy
    2: RGBColor(0x2E, 0x6D, 0xA4),   # medium blue
    3: RGBColor(0x2E, 0x86, 0xAB),   # steel blue
    4: RGBColor(0x45, 0x8B, 0x74),   # teal green
}

def set_heading(text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.color.rgb = HEADING_COLORS.get(level, RGBColor(0, 0, 0))
    run.font.bold = True
    run.font.size = Pt({1: 18, 2: 14, 3: 12, 4: 11}.get(level, 11))
    return h

def shade_cell(cell, hex_color='D9EAF7'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], 'D9EAF7')
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                shade_cell(cells[ci], 'F5FAFF')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return t

def add_code_block(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F0F0F0')
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_note(text, color=RGBColor(0x44, 0x6E, 0x91), icon='ℹ'):
    p = doc.add_paragraph()
    run = p.add_run(icon + '  ' + text)
    run.font.italic = True
    run.font.size   = Pt(10)
    run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.2)
    return p

def add_warning(text):
    return add_note(text, color=RGBColor(0xA6, 0x3D, 0x2E), icon='⚠')

def add_success(text):
    return add_note(text, color=RGBColor(0x2E, 0x7D, 0x4F), icon='✔')

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    text = re.sub(r'^[-*]\s*', '', text)
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return p

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    return p

def add_para(text):
    p = doc.add_paragraph()
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.bold = (i % 2 == 1)
        run.font.size = Pt(10.5)
    return p

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Inventory In/Out Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
title.runs[0].font.size = Pt(26)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Investigation & Fix Documentation')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = meta.add_run('Tri-E Enterprises · TOS System · Reports Module')
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = meta2.add_run('Prepared: August 22, 2026')
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════════
set_heading('Executive Summary', 1)
add_para('The client reported that the **Inventory In/Out Report** was not showing stock-in transactions — specifically, a staff member said they entered **700 pcs of APO Cement (Portland)** but it never appeared in the report.')
doc.add_paragraph()
add_para('Investigation uncovered **three separate defects** in how the system tracks stock movements. None of them were visible from the report screen itself — each required tracing the data through the Purchases, Products, and Audit Log modules to find. All three have been fixed.')
doc.add_paragraph()

add_table(
    ['#', 'Issue Found', 'Severity', 'Status'],
    [
        ['1', 'Purchases (stock received from suppliers) were never recorded as "in" movements', 'High — hid the majority of real stock-in activity', 'Fixed + backfilled'],
        ['2', 'The "Stock" field on Edit Product silently overwrites inventory with no audit trail', 'Critical — no record of who changed stock or by how much', 'Fixed'],
        ['3', 'Any staff role with product-edit rights could freely overwrite live stock counts', 'High — data integrity / accountability risk', 'Fixed'],
    ],
    [0.4, 3.6, 2.2, 1.3]
)

add_note('The specific 700 pcs the client asked about was traced to Issue 2 below — it was real and did change the system\'s stock count, it just was never logged as a movement, so the report never showed it.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# ISSUE 1
# ════════════════════════════════════════════════════════════════════════════
set_heading('Issue 1 — Purchases Never Logged as "In" Movements', 1)

set_heading('What was happening', 2)
add_para('The Inventory In/Out Report reads from a dedicated `inventory_movements` table. Before this fix, only two things ever wrote to it:')
for b in ['A sale (POS checkout) — logged as "out"', 'A voided sale — logged as "in" (reversal)']:
    add_bullet(b)
add_para('**Receiving stock through Purchases — the main way stock is supposed to come in — was never logged at all.** The `PurchaseItem` model correctly increased the real stock count (`inventories.quantity`) when a delivery was received, but it never created a matching movement record.')
doc.add_paragraph()

set_heading('Evidence', 2)
add_table(
    ['Metric', 'Before Fix'],
    [
        ['Total "in" movements in the system', '35 (100% from voided sales)'],
        ['Actual purchase receipts on record', '108'],
        ['Purchase receipts visible in the report', '0'],
    ],
    [3.5, 3.0]
)

set_heading('Root Cause', 2)
add_para('File: **app/Models/PurchaseItem.php**')
add_para('The `created`, `updated`, and `deleted` lifecycle hooks updated `Inventory.quantity` directly but never called `InventoryMovement::create()`.')

set_heading('Fix Applied', 2)
for b in [
    'Added `InventoryMovement` logging to all three lifecycle hooks in `PurchaseItem.php`:',
]:
    add_bullet(b)
add_bullet('**Created** (new receipt) → logs an "in" movement, reason "Purchase received"', level=1)
add_bullet('**Updated** (received quantity corrected) → logs the difference as "in" or "out", reason "Purchase receipt adjusted"', level=1)
add_bullet('**Deleted** (receipt line removed) → logs an "out" movement, reason "Purchase item deleted"', level=1)
doc.add_paragraph()
add_para('Every logged movement references the originating Purchase (`reference_type` / `reference_id`) so it can be traced back.')

set_heading('Historical Data Backfill', 2)
add_para('The 108 existing purchase receipts that predated this fix would still be invisible in past reports. With the client\'s approval, these were backfilled:')
for b in [
    '108 historical "in" movement records created, one per past purchase receipt',
    'Each backfilled record was dated to match its **original** purchase timestamp — not today\'s date — so historical reports read correctly',
]:
    add_bullet(b)
add_warning('First backfill attempt used the wrong date (today\'s date on all 108 records) because the `InventoryMovement` model\'s mass-assignment protection silently dropped the explicit `created_at` value. This was caught during verification and corrected by re-running the backfill with a direct database insert that preserves the original timestamps.')

add_success('Result: "in" movements in the system went from 35 → 143. All 108 backfilled records now carry their correct historical dates.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# ISSUE 2
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
set_heading('Issue 2 — The "Stock" Field Silently Overwrites Inventory', 1)

set_heading('The Client\'s Report: 700 pcs of APO Cement', 2)
add_para('The client said a staff member entered 700 pcs of APO Cement (Portland), but it was not visible anywhere in the Inventory In/Out Report.')

set_heading('Investigation Trail', 2)
add_numbered('Searched for the product — found **APO CEMENT (PORTLAND)**, product #633.')
add_numbered('Checked its Purchase history — only one purchase on record (100 pcs, March 17, 2026). No 700 pcs purchase existed.')
add_numbered('Checked its Inventory Movement history — no "in" entry of 700 pcs anywhere.')
add_numbered('Checked other cement products in case of a mix-up (Mayon Cement, Grand Cement T1P) — nothing matching there either.')
add_numbered('The client then supplied a screenshot of the **Audit Log**, which showed the real event:')
doc.add_paragraph()

add_table(
    ['Field', 'Value'],
    [
        ['User', 'Warehouse'],
        ['Action', 'Updated'],
        ['Record', 'Product: APO CEMENT (PORTLAND)'],
        ['Old Value (quantity)', '2'],
        ['New Value (quantity)', '700'],
        ['Timestamp', 'Aug 21, 2026, 12:42:34 PM'],
    ],
    [2.5, 4.0]
)

set_heading('Root Cause', 2)
add_para('File: **app/Filament/Resources/Products/Pages/EditProduct.php**')
add_para('Every Product record has an editable **"Stock"** field on its Edit screen. On save, this ran:')
add_code_block("protected function afterSave(): void\n{\n    $quantity = $this->data['quantity'] ?? 0;\n\n    $this->record->inventory()->updateOrCreate(\n        ['product_id' => $this->record->id],\n        ['quantity' => $quantity]   // <- overwrites, no movement logged\n    );\n}")
add_para('This **directly overwrote** the real, working inventory count with whatever number was typed — it does not add to stock, it replaces it — and it never created an `InventoryMovement` record. So the change was completely real (it changed the actual stock number the POS uses) but totally invisible in the audit trail the report relies on.')

set_heading('Confirming the Numbers Match', 2)
add_para('The math confirms this is exactly what happened:')
add_table(
    ['Event', 'Quantity'],
    [
        ['Stock set via Edit Product, Aug 21 12:42 PM', '700'],
        ['Pcs sold via POS since then (10 sales)', '− 272'],
        ['Current on-hand quantity (verified in system)', '= 428'],
    ],
    [4.5, 2.0]
)
add_success('428 matches the system\'s actual current stock exactly — confirming the 700 pcs entry was real and did take effect, it simply left no trace in the report.')

set_heading('Fix Applied', 2)
add_para('`EditProduct.php` now calculates the **difference** between the old and new quantity on every save, and logs it as a proper `InventoryMovement`:')
for b in [
    'Positive difference → logged as "in"',
    'Negative difference → logged as "out"',
    'Reason recorded as **"Manual Stock Adjustment"**',
    'Notes capture **who** made the change and the **old → new** values, e.g. "Stock edited via Product form by Warehouse (2 → 700)"',
]:
    add_bullet(b)
doc.add_paragraph()
add_para('The same logging was added to **CreateProduct.php**, so setting an opening stock quantity when adding a brand-new product is now logged too, as "Initial Stock".')
doc.add_paragraph()
add_note('Per the client\'s decision, the Aug 21 event itself was **not** backfilled into historical reports — this fix applies to all edits going forward only.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# ISSUE 3
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
set_heading('Issue 3 — Access Control on the Stock Field', 1)

set_heading('The Problem', 2)
add_para('Even with logging now in place, the underlying design was still risky: **any role with product-edit permission could freely overwrite live stock counts** from a plain number field, with no reason required and no protection against accidental overwrites (e.g. mistyping "700" when meaning to add 7, or editing a stale number).')
doc.add_paragraph()
add_para('The "Warehouse clerk" role — meant for receiving stock — had `Update:Product` permission, which is how the field became accessible for this kind of edit in the first place.')

set_heading('Fix Applied', 2)
add_para('File: **app/Filament/Resources/Products/Schemas/ProductForm.php**')
add_para('On the **Edit Product** screen, the "Stock" field is now **read-only** for every role except **super_admin** and **Ops Sup** (Operation Supervisor). Restricted users still see the current stock number, with a note explaining the alternative:')
add_code_block('"Only admins can adjust stock here. Use Purchases to receive new stock."')
add_para('The **Create Product** screen is unaffected — setting an opening stock quantity for a brand-new product is still available to anyone who can create products, and is now properly logged (see Issue 2).')

set_heading('Resulting Workflow by Role', 2)
add_table(
    ['Role', 'Receive Stock (Purchases)', 'Overwrite Stock Field (Edit Product)'],
    [
        ['super_admin', 'Yes', 'Yes'],
        ['Ops Sup', 'Yes', 'Yes'],
        ['Warehouse clerk', 'Yes', 'No — read only'],
        ['cashier', 'N/A', 'No — read only'],
        ['Sales Rep', 'N/A', 'No — read only'],
        ['Driver', 'N/A', 'No — read only'],
    ],
    [2.2, 2.4, 2.9]
)
add_note('Staff now have exactly one clear, fully-audited path to add stock: the Purchases module.')
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SUMMARY OF FILES CHANGED
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
set_heading('Summary of Files Changed', 1)
add_table(
    ['File', 'Change'],
    [
        ['app/Models/PurchaseItem.php',
         'Logs an InventoryMovement whenever a purchase receipt is created, its received quantity is adjusted, or it is deleted.'],
        ['app/Filament/Resources/Products/Pages/EditProduct.php',
         'Logs the stock quantity difference as an InventoryMovement ("Manual Stock Adjustment") on every product save.'],
        ['app/Filament/Resources/Products/Pages/CreateProduct.php',
         'Logs an InventoryMovement ("Initial Stock") when a new product is created with an opening stock quantity.'],
        ['app/Filament/Resources/Products/Schemas/ProductForm.php',
         'Restricts the Stock field to read-only on Edit Product for all roles except super_admin and Ops Sup.'],
    ],
    [3.0, 4.0]
)

set_heading('Data Changes Applied', 2)
for b in [
    '**108** historical InventoryMovement "in" records were created for past purchase receipts, each dated to match its original purchase timestamp.',
    'No historical data was altered for Issue 2 (the Aug 21 manual stock edit) — only the code was fixed, per the client\'s decision.',
]:
    add_bullet(b)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# RECOMMENDATIONS
# ════════════════════════════════════════════════════════════════════════════
set_heading('Recommended Next Step (Not Yet Implemented)', 1)
add_para('Even restricted to admins, the Stock field on Edit Product is still an **absolute overwrite** — typing a number replaces the count rather than adjusting it, which carries some risk of an accidental miskey erasing the true stock level.')
doc.add_paragraph()
add_para('Recommendation: replace it with a proper **Stock Adjustment** action:')
for b in [
    'Delta-based entry (e.g. "+50" or "−12") instead of an absolute number',
    'A **required reason** — e.g. Damaged, Miscount / Stock Take Correction, Other',
    'Still restricted to admin / Ops Sup roles',
    'Logged the same way as the other movement types, for full consistency',
]:
    add_bullet(b)
doc.add_paragraph()
add_note('This was proposed to the client and is pending a decision on whether to proceed.')

# ════════════════════════════════════════════════════════════════════════════
doc.save('Tri-E Enterprises - Inventory Report Fix Documentation.docx')
print('Saved.')
